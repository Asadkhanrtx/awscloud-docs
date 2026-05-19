import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style helpers
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        set_font(run, size=20, bold=True, color=(15, 52, 96))
    elif level == 2:
        set_font(run, size=15, bold=True, color=(22, 33, 62))
    elif level == 3:
        set_font(run, size=12, bold=True, color=(15, 52, 96))
    elif level == 4:
        set_font(run, size=11, bold=True, color=(83, 52, 131))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_paragraph(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    # Handle inline bold/italic/code
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, color=(15, 52, 96))
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=10, color=(192, 57, 43))
            run.font.highlight_color = None
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_font(run)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # shade background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A1A2E')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=9, color=(220, 220, 220))
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0F3460')
        tcPr.append(shd)
    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, size=10, bold=True, color=(15, 52, 96))
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    set_font(run, name="Courier New", size=9, color=(192, 57, 43))
                else:
                    run = p.add_run(part)
                    set_font(run, size=10)
            if r_idx % 2 == 1:
                tc = tr.cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F4FF')
                tcPr.append(shd)
    doc.add_paragraph()

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    p.paragraph_format.space_after = Pt(3)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, color=(15, 52, 96))
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=10, color=(192, 57, 43))
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_font(run)

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4FF')
    pPr.append(shd)
    run = p.add_run(text.lstrip('> '))
    set_font(run, italic=True, color=(70, 70, 120))

# ── Parse markdown ──────────────────────────────────────────
with open("vpc-case-study.md", "r") as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip image lines
    if re.match(r'!\[.*?\]\(.*?\)', line):
        i += 1
        continue

    # Headings
    if line.startswith('#### '):
        add_heading(line[5:], 4); i += 1; continue
    if line.startswith('### '):
        add_heading(line[4:], 3); i += 1; continue
    if line.startswith('## '):
        add_heading(line[3:], 2); i += 1; continue
    if line.startswith('# '):
        add_heading(line[2:], 1); i += 1; continue

    # HR
    if re.match(r'^---+$', line):
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        i += 1; continue

    # Blockquote
    if line.startswith('> '):
        add_blockquote(line); i += 1; continue

    # Fenced code block
    if line.startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].rstrip('\n').startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        add_code_block('\n'.join(code_lines))
        i += 1; continue

    # Table
    if '|' in line and i + 1 < len(lines) and re.match(r'[\|\- ]+', lines[i+1]):
        table_lines = []
        while i < len(lines) and '|' in lines[i]:
            table_lines.append(lines[i].rstrip('\n'))
            i += 1
        # Parse headers and rows
        raw_headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
        raw_rows = []
        for tl in table_lines[2:]:
            row = [c.strip() for c in tl.split('|') if c.strip()]
            if row:
                raw_rows.append(row)
        if raw_headers:
            add_table(raw_headers, raw_rows)
        continue

    # Bullet list
    if re.match(r'^[-*] ', line):
        add_bullet(line[2:]); i += 1; continue
    if re.match(r'^\d+\. ', line):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(3)
        text = re.sub(r'^\d+\. ', '', line)
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_font(run, bold=True, color=(15, 52, 96))
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                set_font(run, name="Courier New", size=10, color=(192, 57, 43))
            else:
                run = p.add_run(part)
                set_font(run)
        i += 1; continue

    # Empty line
    if line.strip() == '':
        i += 1; continue

    # Normal paragraph
    add_paragraph(line)
    i += 1

doc.save("vpc-case-study.docx")
print("Word document created: vpc-case-study.docx")
