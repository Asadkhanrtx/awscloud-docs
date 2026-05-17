from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for sec in doc.sections:
    sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.8)
    sec.left_margin=Cm(2.2); sec.right_margin=Cm(2.2)

NAVY  =RGBColor(0x0F,0x34,0x60); DARK =RGBColor(0x16,0x21,0x3E)
WHITE =RGBColor(0xFF,0xFF,0xFF); RED  =RGBColor(0xE9,0x45,0x60)
PURP  =RGBColor(0x53,0x34,0x83); GREY =RGBColor(0x44,0x44,0x44)
GOLD  =RGBColor(0x85,0x64,0x04); LBLUE=RGBColor(0x0F,0x34,0x60)
GREEN =RGBColor(0x1E,0x8B,0x4C); ORAN =RGBColor(0xE6,0x7E,0x22)

def cshd(cell,hex):
    tc=cell._tc;p=tc.get_or_add_tcPr()
    for o in p.findall(qn('w:shd')):p.remove(o)
    s=OxmlElement('w:shd');s.set(qn('w:val'),'clear')
    s.set(qn('w:color'),'auto');s.set(qn('w:fill'),hex);p.append(s)

def pshd(p,hex):
    pp=p._p.get_or_add_pPr()
    for o in pp.findall(qn('w:shd')):pp.remove(o)
    s=OxmlElement('w:shd');s.set(qn('w:val'),'clear')
    s.set(qn('w:color'),'auto');s.set(qn('w:fill'),hex);pp.append(s)

def pbdr_left(p,color='0F3460',sz='18'):
    pp=p._p.get_or_add_pPr();b=OxmlElement('w:pBdr')
    l=OxmlElement('w:left');l.set(qn('w:val'),'single')
    l.set(qn('w:sz'),sz);l.set(qn('w:space'),'4');l.set(qn('w:color'),color)
    b.append(l);pp.append(b)

def pbdr_bot(p,color='CCCCCC',sz='6'):
    pp=p._p.get_or_add_pPr();b=OxmlElement('w:pBdr')
    e=OxmlElement('w:bottom');e.set(qn('w:val'),'single')
    e.set(qn('w:sz'),sz);e.set(qn('w:space'),'1');e.set(qn('w:color'),color)
    b.append(e);pp.append(b)

def sp(p,b=0,a=4):
    p.paragraph_format.space_before=Pt(b);p.paragraph_format.space_after=Pt(a)

def rfmt(r,name='Calibri',size=11,bold=False,italic=False,color=None):
    r.font.name=name;r.font.size=Pt(size)
    r.font.bold=bold;r.font.italic=italic
    if color:r.font.color.rgb=color

def ir(p,text,name='Calibri',size=11,bold=False,italic=False,color=None):
    r=p.add_run(text);rfmt(r,name,size,bold,italic,color);return r

def wri(p,text,size=11):
    import re
    for part in re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)',text):
        if part.startswith('**') and part.endswith('**'):
            ir(p,part[2:-2],size=size,bold=True,color=NAVY)
        elif part.startswith('`') and part.endswith('`'):
            ir(p,part[1:-1],name='Courier New',size=size-1,color=RGBColor(0xC0,0x39,0x2B))
        elif part.startswith('*') and part.endswith('*'):
            ir(p,part[1:-1],size=size,italic=True,color=GREY)
        else:
            ir(p,part,size=size)

def h1(text):
    p=doc.add_paragraph();sp(p,b=12,a=6);pbdr_bot(p,'E94560','12')
    r=p.add_run(text);rfmt(r,size=20,bold=True,color=NAVY);return p

def h2(text):
    p=doc.add_paragraph();sp(p,b=10,a=4)
    r=p.add_run('  '+text);rfmt(r,size=14,bold=True,color=DARK)
    pbdr_left(p,'0F3460','18');return p

def h3(text):
    p=doc.add_paragraph();sp(p,b=8,a=3)
    r=p.add_run(text);rfmt(r,size=12,bold=True,color=NAVY);return p

def h4(text):
    p=doc.add_paragraph();sp(p,b=6,a=2)
    r=p.add_run(text);rfmt(r,size=11,bold=True,color=PURP);return p

def body(text,b=0,a=5):
    p=doc.add_paragraph();sp(p,b=b,a=a);wri(p,text,size=11);return p

def bul(text,level=0,b=1,a=2):
    p=doc.add_paragraph(style='List Bullet');sp(p,b=b,a=a)
    p.paragraph_format.left_indent=Inches(0.25+level*0.2)
    wri(p,text,size=11);return p

def num(text,b=1,a=2):
    p=doc.add_paragraph(style='List Number');sp(p,b=b,a=a)
    p.paragraph_format.left_indent=Inches(0.25);wri(p,text,size=11);return p

def cod(text,b=4,a=4):
    p=doc.add_paragraph();sp(p,b=b,a=a)
    p.paragraph_format.left_indent=Inches(0.2)
    p.paragraph_format.right_indent=Inches(0.2)
    pshd(p,'1A1A2E')
    r=p.add_run(text);rfmt(r,name='Courier New',size=9,color=RGBColor(0xDC,0xDC,0xDC));return p

def note(text,b=4,a=4):
    p=doc.add_paragraph();sp(p,b=b,a=a)
    p.paragraph_format.left_indent=Inches(0.25)
    pshd(p,'FFF3CD')
    r=p.add_run('  ⚠  '+text);rfmt(r,size=10,italic=True,color=GOLD);return p

def info(text,b=4,a=4):
    p=doc.add_paragraph();sp(p,b=b,a=a)
    p.paragraph_format.left_indent=Inches(0.25)
    pshd(p,'E8F4FD')
    r=p.add_run('  ℹ  '+text);rfmt(r,size=10,color=LBLUE);return p

def div():
    p=doc.add_paragraph();sp(p,b=4,a=4);pbdr_bot(p,'DDDDDD','4');return p

def tbl(headers,rows,col_widths=None):
    t=doc.add_table(rows=1+len(rows),cols=len(headers))
    t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text=''
        p=c.paragraphs[0];sp(p,b=2,a=2)
        r=p.add_run(h);rfmt(r,size=10,bold=True,color=WHITE)
        cshd(c,'0F3460');c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    for ri,row in enumerate(rows):
        fill='F0F4FF' if ri%2==1 else 'FFFFFF'
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci];c.text=''
            p=c.paragraphs[0];sp(p,b=2,a=2)
            wri(p,str(val),size=10);cshd(c,fill)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows:row.cells[i].width=Inches(w)
    p=doc.add_paragraph();sp(p,b=0,a=2);return t

def arch(lines,title=None):
    if title:
        p=doc.add_paragraph();sp(p,b=6,a=0);pshd(p,'533483')
        r=p.add_run(f'  Architecture — {title}')
        rfmt(r,size=10,bold=True,color=WHITE)
    cod('\n'.join(lines),b=0,a=6)

def step_banner(n,title,color='16213E'):
    p=doc.add_paragraph();sp(p,b=8,a=4);pshd(p,color)
    r1=p.add_run(f'  Task {n} — ');rfmt(r1,size=12,bold=True,color=RED)
    r2=p.add_run(title);rfmt(r2,size=12,bold=True,color=WHITE)

def section_banner(title,sub=''):
    p=doc.add_paragraph();sp(p,b=0,a=0);pshd(p,'0F3460')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title);rfmt(r,size=16,bold=True,color=WHITE)
    if sub:
        p2=doc.add_paragraph();sp(p2,b=0,a=8);pshd(p2,'16213E')
        p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r2=p2.add_run(sub);rfmt(r2,size=10,color=RGBColor(0xAA,0xCC,0xFF))

# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════
p=doc.add_paragraph();sp(p,b=0,a=0);pshd(p,'0F3460')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
ir(p,'AWS LOAD BALANCER',size=26,bold=True,color=WHITE)

p=doc.add_paragraph();sp(p,b=0,a=0);pshd(p,'16213E')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
ir(p,'Complete Case Study Documentation  ·  Module: Elastic Load Balancing',size=11,color=RGBColor(0xAA,0xCC,0xFF))

p=doc.add_paragraph();sp(p,b=0,a=10);pshd(p,'E94560')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
ir(p,'ALB  |  NLB  |  GLB  |  Path Routing  |  Host Routing  |  Query Routing  |  HTTP Method Routing  |  NLB→ALB',size=10,bold=True,color=WHITE)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT IS A LOAD BALANCER
# ═══════════════════════════════════════════════════════════════════════════════
h1('1. What is a Load Balancer?')
h3('Definition')
body('A Load Balancer is a networking service that distributes incoming traffic across multiple targets (EC2 instances, containers, IP addresses, Lambda functions) in one or more Availability Zones to ensure:')
bul('High Availability'); bul('Fault Tolerance'); bul('Scalability'); bul('Reliability')

div()
h3('Real-World Analogy')
p=doc.add_paragraph();sp(p,b=0,a=0);pshd(p,'FFF3CD')
ir(p,'  Imagine a restaurant with multiple billing counters.',size=11,italic=True,color=GOLD)
p=doc.add_paragraph();sp(p,b=0,a=0);pshd(p,'FFF3CD')
ir(p,'  A manager stands at the entrance and sends customers to the counter with the ',size=11,italic=True,color=GOLD)
ir(p,'shortest queue',size=11,bold=True,color=RGBColor(0x8B,0x45,0x13))
ir(p,'.',size=11,italic=True,color=GOLD)
p=doc.add_paragraph();sp(p,b=0,a=6);pshd(p,'FFF3CD')
ir(p,'  That manager is exactly like a Load Balancer.',size=11,bold=True,color=RGBColor(0x8B,0x45,0x13))

body('**Result:**')
bul('No single counter gets overloaded'); bul('Customers are served faster')
bul('If one counter fails, others continue working')

arch([
    '  Customers (Traffic)',
    '         |',
    '  [Manager = Load Balancer]  ← single entry point',
    '   /        |        \\',
    '[Counter1] [Counter2] [Counter3]   ← Backend Servers',
    '',
    '  If Counter2 breaks → Manager stops sending customers there (health check)',
    '  New counter added → Manager starts sending customers automatically (Auto Scaling)',
],'Restaurant = Load Balancer Analogy')

div()
h3('Key Benefits of Load Balancers')
tbl(['#','Benefit','Description'],[
    ['1','**High Availability**','Distributes traffic across multiple Availability Zones'],
    ['2','**Fault Tolerance**','Stops sending traffic to unhealthy servers'],
    ['3','**Scalability**','Handles increasing traffic automatically'],
    ['4','**Security**','Backend servers are hidden behind the LB'],
    ['5','**Health Checks**','Continuously monitors target health'],
    ['6','**SSL/TLS Termination**','Offloads HTTPS encryption/decryption'],
    ['7','**Better Performance**','Prevents server overload'],
],col_widths=[0.3,2.0,4.2])

div()
h3('AWS Elastic Load Balancing (ELB) Family')
arch([
    '  ┌──────────────────────────────────────────────────────────────┐',
    '  │              AWS Elastic Load Balancing (ELB)               │',
    '  ├──────────────────────┬──────────────────────┬───────────────┤',
    '  │         ALB          │         NLB          │      GLB      │',
    '  │  Application LB      │    Network LB        │  Gateway LB   │',
    '  │      Layer 7         │      Layer 4         │    Layer 3    │',
    '  │    HTTP/HTTPS        │    TCP/UDP/TLS        │   IP Traffic  │',
    '  └──────────────────────┴──────────────────────┴───────────────┘',
],'ELB Family')
note('Classic Load Balancer (CLB) is legacy/deprecated. AWS recommends ALB or NLB for new projects.')

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ALB
# ═══════════════════════════════════════════════════════════════════════════════
div()
h1('2. Application Load Balancer (ALB) — Layer 7')
h3('Definition')
body('An **Application Load Balancer (ALB)** operates at **OSI Layer 7 (Application Layer)**.')
body('It makes routing decisions based on the **content of HTTP/HTTPS requests** such as:')
bul('URL Path'); bul('Hostname'); bul('HTTP Headers'); bul('Query Strings'); bul('HTTP Methods')

div()
h3('ALB Key Characteristics')
tbl(['Feature','Description'],[
    ['Protocol Support','HTTP, HTTPS, gRPC, WebSocket'],
    ['OSI Layer','Layer 7 (Application Layer)'],
    ['Routing Type','Content-Based Routing'],
    ['Targets Supported','EC2, ECS, Lambda, IP addresses'],
    ['SSL Termination','Supported'],
    ['Sticky Sessions','Supported'],
    ['Health Checks','HTTP/HTTPS path-based'],
    ['Cross-Zone LB','Enabled by default (free)'],
],col_widths=[2.5,4.0])

div()
h3('ALB Architecture')
arch([
    '                     ┌────────────────────────────────┐',
    '                     │     Application Load Balancer  │',
    'Client ──HTTP/HTTPS─►│             (ALB)              │',
    '                     │  Listener Rules                │',
    '                     │                                │',
    '                     │  IF path = /api/*      ─────►  API Target Group',
    '                     │  IF path = /images/*   ─────►  Image Target Group',
    '                     │  IF host = mobile.*    ─────►  Mobile Target Group',
    '                     │  DEFAULT               ─────►  Web Target Group',
    '                     └────────────────────────────────┘',
],'ALB Content-Based Routing')

div()
h3('ALB Core Components')
tbl(['Component','Description'],[
    ['**Listener**','Checks incoming requests on a specific port'],
    ['**Listener Rule**','Defines how traffic should be routed'],
    ['**Target Group**','Group of backend servers/targets'],
    ['**Health Check**','Monitors target health status'],
],col_widths=[2.0,4.5])

div()
h3('ALB Listener Rules — One-Line Definitions')
tbl(['Rule Type','One-Line Definition'],[
    ['**Path-Based Routing**','Routes traffic based on URL path'],
    ['**Host-Based Routing**','Routes traffic based on domain/hostname'],
    ['**Header-Based Routing**','Routes traffic based on HTTP headers'],
    ['**Query String Routing**','Routes traffic using query parameters'],
    ['**HTTP Method Routing**','Routes based on GET, POST, PUT, DELETE methods'],
    ['**Source IP Routing**','Routes traffic based on client IP range'],
],col_widths=[2.5,4.0])

div()
h3('ALB Routing Examples')

h4('1. Path-Based Routing')
cod('example.com/api/*      ───► API Servers\nexample.com/images/*   ───► Image Servers\nexample.com/admin/*    ───► Admin Servers')

h4('2. Host-Based Routing')
cod('api.example.com        ───► API Target Group\nwww.example.com        ───► Web Target Group\nadmin.example.com      ───► Admin Target Group')

h4('3. Header-Based Routing')
cod('Header: Device=Mobile  ───► Mobile Backend\nHeader: Device=Web     ───► Web Backend')

h4('4. Query String Routing')
cod('?platform=mobile       ───► Mobile Servers\n?version=v2            ───► Version 2 Backend')

h4('5. HTTP Method Routing')
cod('GET Requests           ───► Read Servers\nPOST Requests          ───► Write Servers')

h4('6. Source IP Routing')
cod('Corporate IP Range     ───► Internal Application\nPublic Users           ───► Public Backend')

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — INTERNET-FACING vs INTERNAL ALB
# ═══════════════════════════════════════════════════════════════════════════════
div()
h1('3. Internet-Facing ALB vs Internal ALB')

h2('Internet-Facing ALB')
h3('Definition')
body('An **Internet-Facing ALB** has:')
bul('Public IP addresses'); bul('Public DNS name'); bul('Internet accessibility')
body('It receives traffic directly from the internet.')

arch([
    '                           INTERNET',
    '                               │',
    '                               ▼',
    '                ┌───────────────────────────┐',
    '                │   Internet-Facing ALB     │',
    '                │  Public IP / Public DNS   │',
    '                │  Deployed in PUBLIC       │',
    '                │  Subnets                  │',
    '                └────────────┬──────────────┘',
    '                             │',
    '      ┌──────────────────────┼──────────────────────┐',
    '      ▼                      ▼                      ▼',
    '┌──────────┐           ┌──────────┐           ┌──────────┐',
    '│  EC2-1   │           │  EC2-2   │           │  EC2-3   │',
    '│ Private  │           │ Private  │           │ Private  │',
    '└──────────┘           └──────────┘           └──────────┘',
    '       Backend instances can remain PRIVATE',
],'Internet-Facing ALB')

tbl(['Feature','Details'],[
    ['DNS Name','Resolves to Public IPs'],
    ['ALB Placement','Public Subnets'],
    ['Targets','Usually Private Subnets'],
    ['Security Group','Allow 80/443 from 0.0.0.0/0'],
    ['Accessed By','Public Internet Users'],
    ['Use Cases','Websites, Public APIs'],
],col_widths=[2.5,4.0])

div()
h2('Internal ALB')
h3('Definition')
body('An **Internal ALB** has:')
bul('Private IP addresses only'); bul('Private DNS name'); bul('No internet accessibility')
body('Used for internal communication inside a VPC.')

arch([
    '       ┌─────────────────────────┐',
    '       │   Internet-Facing ALB   │',
    '       │     Public Traffic      │',
    '       └──────────┬──────────────┘',
    '                  │',
    '                  ▼',
    '       ┌─────────────────────────┐',
    '       │       Web Tier          │',
    '       │     Frontend EC2        │',
    '       └──────────┬──────────────┘',
    '                  │',
    '                  ▼',
    '       ┌─────────────────────────┐',
    '       │      Internal ALB       │',
    '       │ Private IP / Private DNS│',
    '       │  Deployed in PRIVATE    │',
    '       │        Subnets          │',
    '       └──────────┬──────────────┘',
    '                  │',
    '     ┌────────────┼────────────┐',
    '     ▼            ▼            ▼',
    '┌─────────┐  ┌─────────┐  ┌─────────┐',
    '│ App-1   │  │ App-2   │  │ App-3   │',
    '│ Private │  │ Private │  │ Private │',
    '└─────────┘  └─────────┘  └─────────┘',
],'Internal ALB in Three-Tier Architecture')

tbl(['Feature','Details'],[
    ['DNS Name','Resolves to Private IPs'],
    ['Placement','Private Subnets'],
    ['Access','Only inside VPC'],
    ['Security Group','Allow VPC CIDR or specific SGs'],
    ['Use Cases','Microservices, Internal APIs'],
],col_widths=[2.5,4.0])

div()
h3('Internet-Facing ALB vs Internal ALB — Comparison')
tbl(['Feature','Internet-Facing ALB','Internal ALB'],[
    ['IP Type','Public IP','Private IP'],
    ['DNS Resolution','Public DNS','Private DNS'],
    ['Subnets','Public','Private'],
    ['Accessible From','Internet + VPC','VPC Only'],
    ['Security Group','0.0.0.0/0','VPC CIDR / Specific SG'],
    ['Position','Edge Entry Point','Between Application Tiers'],
    ['Example','www.example.com','order-service.internal'],
],col_widths=[2.0,2.2,2.3])

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — NLB
# ═══════════════════════════════════════════════════════════════════════════════
div()
h1('4. Network Load Balancer (NLB) — Layer 4')
h3('Definition')
body('A **Network Load Balancer (NLB)** operates at **OSI Layer 4 (Transport Layer)**.')
body('It routes traffic based on:')
bul('IP Address'); bul('TCP/UDP Port')
body('It does NOT inspect request content.')
body('Designed for:')
bul('Ultra High Performance'); bul('Very Low Latency'); bul('Millions of requests per second')

div()
h3('NLB Key Characteristics')
tbl(['Feature','Description'],[
    ['Protocol Support','TCP, UDP, TLS'],
    ['OSI Layer','Layer 4'],
    ['Routing Type','IP + Port Based'],
    ['Performance','Millions of requests/sec'],
    ['Latency','Microseconds'],
    ['Static IP','Supported'],
    ['Elastic IP','Supported'],
    ['Preserve Client IP','Yes'],
    ['SSL Termination','Supported'],
    ['Health Checks','TCP/HTTP/HTTPS'],
],col_widths=[2.5,4.0])

div()
h3('NLB Architecture')
arch([
    '                        INTERNET',
    '                            │',
    '                            ▼',
    '                ┌────────────────────┐',
    '                │ Network Load       │',
    '                │ Balancer (NLB)     │',
    '                │ Layer 4            │',
    '                │ Static IP          │',
    '                └─────────┬──────────┘',
    '                          │',
    '                          │ Routes using IP + Port',
    '                          │',
    '      ┌───────────────────┼───────────────────┐',
    '      ▼                   ▼                   ▼',
    ' ┌─────────┐         ┌─────────┐         ┌─────────┐',
    ' │ EC2-1   │         │ EC2-2   │         │ EC2-3   │',
    ' │ :8080   │         │ :8080   │         │ :8080   │',
    ' └─────────┘         └─────────┘         └─────────┘',
],'NLB Layer 4 Routing')

div()
h3('Why Static IP in NLB Matters')
h4('ALB')
cod('my-alb-123.elb.amazonaws.com\n→ Dynamic IPs\n→ IPs can change anytime\n→ Use DNS name only')
h4('NLB')
cod('52.10.20.30\n→ Static IP\n→ Never changes\n→ Can attach Elastic IP\n→ Easy firewall whitelisting')

div()
h3('ALB vs NLB — Detailed Comparison')
tbl(['Feature','ALB','NLB'],[
    ['OSI Layer','Layer 7','Layer 4'],
    ['Protocols','HTTP, HTTPS, gRPC, WebSocket','TCP, UDP, TLS'],
    ['Routing','Content-Based','IP + Port'],
    ['Performance','High','Extreme'],
    ['Latency','Milliseconds','Microseconds'],
    ['Static IP','❌ No','✅ Yes'],
    ['Elastic IP','❌ No','✅ Yes'],
    ['Preserve Client IP','Via X-Forwarded-For','Native'],
    ['SSL Termination','Supported','Supported'],
    ['Sticky Sessions','Cookie-Based','Source IP Based'],
    ['Lambda Support','✅ Yes','❌ No'],
    ['ALB as Target','❌','✅'],
    ['Cross-Zone LB','Enabled by Default','Disabled by Default'],
    ['Security Groups','Supported','Not Supported'],
],col_widths=[2.2,2.2,2.1])

note('NLB does NOT have Security Groups. The backend EC2 instance sees the real client IP. EC2 Security Group must allow client IP ranges directly — traffic does NOT appear from NLB IPs.')

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — PRACTICE TASKS ALB + NLB
# ═══════════════════════════════════════════════════════════════════════════════
div()
section_banner('AWS Load Balancer Practice Tasks (ALB + NLB)','Path-Based Routing · NLB · NLB in Front of ALB')

div()
h3('Architecture Overview')
arch([
    '                    INTERNET',
    '                        │',
    '        ┌───────────────┴───────────────┐',
    '        │                               │',
    '        ▼                               ▼',
    '┌─────────────────┐         ┌─────────────────┐',
    '│       ALB       │         │       NLB       │',
    '│ Path Routing    │         │ TCP Load Bal.   │',
    '└────────┬────────┘         └────────┬────────┘',
    '         │                            │',
    ' ┌───────┴────────┐         ┌─────────┴────────┐',
    ' ▼                ▼         ▼                  ▼',
    'EC2-App1      EC2-App2   EC2-NLB-1        EC2-NLB-2',
    '/app1         /app2      Red Page         Purple Page',
],'ALB + NLB Practice Architecture')

div()
# TASK 1
step_banner(1,'Launch EC2-App1 Instance')
body('**Objective:** Create an EC2 instance that serves the `/app1` application.')
tbl(['Setting','Value'],[
    ['Name','EC2-App1'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['VPC','LB-Practice-VPC'],['Subnet','Public-Subnet-1'],
    ['Public IP','Enabled'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data Script')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
mkdir -p /var/www/html/app1
cat <<EOF > /var/www/html/app1/index.html
<html>
<body style='background-color:#4CAF50; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>Welcome to APP 1 🟢</h1>
<h2 style='color:white;'>This is EC2 - App1 Server</h2>
<h3 style='color:white;'>Path: /app1</h3>
</body>
</html>
EOF''')

div()
# TASK 2
step_banner(2,'Launch EC2-App2 Instance')
body('**Objective:** Create another EC2 instance that serves the `/app2` application.')
tbl(['Setting','Value'],[
    ['Name','EC2-App2'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['Subnet','Public-Subnet-2'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data Script')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
mkdir -p /var/www/html/app2
cat <<EOF > /var/www/html/app2/index.html
<html>
<body style='background-color:#2196F3; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>Welcome to APP 2 🔵</h1>
<h2 style='color:white;'>This is EC2 - App2 Server</h2>
<h3 style='color:white;'>Path: /app2</h3>
</body>
</html>
EOF''')

div()
# TASK 3
step_banner(3,'Launch EC2-NLB-1 Instance')
body('**Objective:** Create backend server 1 for Network Load Balancer testing.')
tbl(['Setting','Value'],[
    ['Name','EC2-NLB-1'],['AMI','Ubuntu'],['Subnet','Public-Subnet-1'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data Script')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
cat <<EOF > /var/www/html/index.html
<html>
<body style='background-color:#FF5722; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>NLB Server 1 🔴</h1>
<h2 style='color:white;'>This request was handled by NLB-Server-1</h2>
</body>
</html>
EOF''')

div()
# TASK 4
step_banner(4,'Launch EC2-NLB-2 Instance')
body('**Objective:** Create backend server 2 for Network Load Balancer testing.')
tbl(['Setting','Value'],[
    ['Name','EC2-NLB-2'],['AMI','Ubuntu'],['Subnet','Public-Subnet-2'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data Script')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
cat <<EOF > /var/www/html/index.html
<html>
<body style='background-color:#9C27B0; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>NLB Server 2 🟣</h1>
<h2 style='color:white;'>This request was handled by NLB-Server-2</h2>
</body>
</html>
EOF''')

div()
# TASK 5
step_banner(5,'Create Target Group for App1')
body('**Objective:** Create a target group for the App1 server.')
tbl(['Setting','Value'],[
    ['Target Group Name','TG-App1'],['Protocol','HTTP'],['Port','80'],
    ['Target Type','Instances'],['Health Check Path','/app1/index.html'],
    ['Registered Target','EC2-App1'],
],col_widths=[2.5,4.0])

# TASK 6
step_banner(6,'Create Target Group for App2')
body('**Objective:** Create a target group for the App2 server.')
tbl(['Setting','Value'],[
    ['Target Group Name','TG-App2'],['Protocol','HTTP'],['Port','80'],
    ['Target Type','Instances'],['Health Check Path','/app2/index.html'],
    ['Registered Target','EC2-App2'],
],col_widths=[2.5,4.0])

div()
# TASK 7
step_banner(7,'Create Application Load Balancer (ALB)')
body('**Objective:** Create an Internet-Facing ALB for path-based routing.')
tbl(['Setting','Value'],[
    ['Name','My-ALB'],['Type','Application Load Balancer'],
    ['Scheme','Internet-Facing'],['Protocol','HTTP'],['Port','80'],
    ['VPC','LB-Practice-VPC'],['Subnets','Public-Subnet-1 & Public-Subnet-2'],
    ['Security Group','LB-SG'],['Default Action','Forward traffic to TG-App1'],
],col_widths=[2.5,4.0])

div()
# TASK 8
step_banner(8,'Configure Path-Based Routing')
body('**Objective:** Route requests to different target groups based on URL paths.')
h4('Rule 1 — App1 Routing')
tbl(['Setting','Value'],[
    ['Condition Type','Path'],['Path Value','/app1*'],
    ['Action','Forward to TG-App1'],['Priority','1'],
],col_widths=[2.5,4.0])
h4('Rule 2 — App2 Routing')
tbl(['Setting','Value'],[
    ['Condition Type','Path'],['Path Value','/app2*'],
    ['Action','Forward to TG-App2'],['Priority','2'],
],col_widths=[2.5,4.0])

arch([
    '  ALB Listener Rules (Path-Based):',
    '',
    '  Request: http://ALB-DNS/app1   → Rule 1 matches → TG-App1 → EC2-App1 🟢',
    '  Request: http://ALB-DNS/app2   → Rule 2 matches → TG-App2 → EC2-App2 🔵',
    '  Request: http://ALB-DNS/other  → Default rule   → TG-App1',
],'Path-Based Routing Rules')

div()
# TASK 9
step_banner(9,'Test ALB Path-Based Routing')
body('**Objective:** Verify that ALB routes traffic correctly.')
cod('http://ALB-DNS/app1\nExpected: Green APP1 page 🟢\n\nhttp://ALB-DNS/app2\nExpected: Blue APP2 page 🔵')

div()
# TASK 10
step_banner(10,'Create Target Group for NLB')
body('**Objective:** Create target group for Network Load Balancer.')
tbl(['Setting','Value'],[
    ['Name','TG-NLB'],['Protocol','TCP'],['Port','80'],
    ['Target Type','Instances'],['Health Check','TCP'],
    ['Registered Targets','EC2-NLB-1, EC2-NLB-2'],
],col_widths=[2.5,4.0])

div()
# TASK 11
step_banner(11,'Create Network Load Balancer (NLB)')
body('**Objective:** Create an Internet-Facing NLB.')
tbl(['Setting','Value'],[
    ['Name','My-NLB'],['Type','Network Load Balancer'],
    ['Scheme','Internet-Facing'],['Protocol','TCP'],['Port','80'],
    ['Subnets','Public-Subnet-1 & Public-Subnet-2'],
    ['Default Action','Forward traffic to TG-NLB'],
],col_widths=[2.5,4.0])

div()
# TASK 12
step_banner(12,'Test NLB')
body('**Objective:** Verify NLB distributes traffic between servers.')
cod('http://NLB-DNS/\n\nSometimes: NLB Server 1 🔴\nSometimes: NLB Server 2 🟣\n\nTraffic gets distributed automatically across both EC2 instances.')

arch([
    '  NLB Traffic Distribution:',
    '',
    '  Request 1 → EC2-NLB-1  🔴 (Red Page)',
    '  Request 2 → EC2-NLB-2  🟣 (Purple Page)',
    '  Request 3 → EC2-NLB-1  🔴',
    '  Request 4 → EC2-NLB-2  🟣',
    '',
    '  NLB uses flow-hash algorithm — same client IP = same server (sticky by default)',
],'NLB Round-Robin Distribution')

div()
# TASK 13
step_banner(13,'Configure NLB in Front of ALB')
body('**Objective:** Configure a **Network Load Balancer (NLB)** to forward traffic to an **Application Load Balancer (ALB)** using the special Application Load Balancer Target Type.')
body('This combines:')
bul('NLB → Static IP + High Performance')
bul('ALB → Smart Layer 7 Routing')

div()
h3('What is "Application Load Balancer" Target Type?')
body('The **Application Load Balancer Target Type** is a special target type available in NLB Target Groups where the NLB can directly forward traffic to an ALB. AWS automatically manages the ALB IP addresses internally.')

h4('Old Method vs New Method')
tbl(['','Old Method ❌','New Method ✅'],[
    ['Flow','NLB → IP Target Group → Manually add ALB IPs','NLB → ALB Target Group → ALB directly'],
    ['Problem/Benefit','ALB IPs change frequently, manual updates needed','AWS manages IP changes automatically'],
    ['Complexity','Needed Lambda automation','Simple setup'],
],col_widths=[1.5,2.5,2.5])

arch([
    '  OLD METHOD (Complex):             NEW METHOD (Simple):',
    '',
    '  NLB                               NLB',
    '   │                                 │',
    '   ▼                                 ▼',
    '  IP Target Group                   ALB Target Group',
    '   │                                 │',
    '   ▼                                 ▼',
    '  Manually add ALB IPs        Application Load Balancer',
    '  (IPs change = broken)        (AWS manages IPs internally)',
],'NLB → ALB: Old vs New Method')

h4('Step 1 — Create Target Group with ALB Type')
body('**Navigation:** EC2 → Target Groups → Create Target Group')
tbl(['Setting','Value'],[
    ['Target Type','Application Load Balancer'],
    ['Name','TG-NLB-to-ALB'],['Protocol','TCP'],['Port','80'],
    ['VPC','LB-Practice-VPC'],['Health Check Protocol','HTTP'],
    ['Health Check Path','/app1/index.html'],['Register Target','My-ALB on Port 80'],
],col_widths=[2.5,4.0])

h4('Step 2 — Update Existing NLB Listener')
body('**Navigation:** EC2 → Load Balancers → My-NLB → Listeners Tab')
num('Select TCP : 80 Listener')
num('Click Actions → Edit Listener')
num('Change Default Action: Forward to TG-NLB-to-ALB')
num('Save Changes')

h4('Step 3 — Test the Architecture')
cod('http://NLB-DNS/app1  →  Expected: GREEN APP1 Page 🟢\nhttp://NLB-DNS/app2  →  Expected: BLUE APP2 Page 🔵')

div()
h3('Complete Request Flow')
cod('''Step 1: User sends request to NLB DNS
            │
            ▼
Step 2: NLB receives TCP traffic
            │
            ▼
Step 3: NLB forwards traffic to ALB
        (AWS internally manages ALB IPs)
            │
            ▼
Step 4: ALB reads HTTP request content
            │
            ▼
Step 5: ALB applies path-based routing
            │
            ├── /app1 → TG-App1 → EC2-App1 🟢
            │
            └── /app2 → TG-App2 → EC2-App2 🔵''')

arch([
    '                     INTERNET',
    '                         │',
    '                         ▼',
    '               ┌─────────────────┐',
    '               │       NLB       │',
    '               │  Static Public  │',
    '               │       IP        │',
    '               └────────┬────────┘',
    '                        │',
    '                        ▼',
    '               ┌─────────────────┐',
    '               │       ALB       │',
    '               │ Smart Routing   │',
    '               └────────┬────────┘',
    '                        │',
    '      ┌─────────────────┼─────────────────┐',
    '      ▼                 ▼                 ▼',
    '    /api              /web             /admin',
    '      │                 │                 │',
    '   API EC2s         Web EC2s         Admin EC2s',
],'NLB → ALB Real World Architecture')

h3('Why This Architecture is Powerful')
tbl(['NLB Feature','Benefit'],[
    ['Static IP','Easy firewall whitelisting'],
    ['High Performance','Handles millions of requests'],
    ['Ultra Low Latency','Fast traffic forwarding'],
],col_widths=[2.5,4.0])
tbl(['ALB Feature','Benefit'],[
    ['Path-Based Routing','Smart traffic distribution'],
    ['Host-Based Routing','Domain-based routing'],
    ['HTTP Awareness','Understands web requests'],
],col_widths=[2.5,4.0])

div()
# TASK 14
step_banner(14,'Summary — What Was Implemented')
cod('''✅ Application Load Balancer (ALB)
✅ Path-Based Routing
✅ Network Load Balancer (NLB)
✅ Multi-Server Traffic Distribution
✅ Apache Web Hosting using User Data
✅ NLB → ALB → EC2 Architecture
✅ Static Public Entry Point
✅ Advanced Routing + High Availability
✅ Scalable Infrastructure
✅ Production-Level AWS Design''')

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — ALB ROUTING METHODS TASKS
# ═══════════════════════════════════════════════════════════════════════════════
div()
section_banner('AWS ALB Routing Methods Practice Tasks','Host-Based · Path-Based · Query String · HTTP Method Routing')

div()
h3('All Routing Methods Overview')
arch([
    '  ALB Listener Rules',
    '         │',
    '  ┌──────┴──────────────────────────────────────────┐',
    '  │  Rule 1: Host Header = app1.example.com → TG-App1',
    '  │  Rule 2: Host Header = app2.example.com → TG-App2',
    '  │  Rule 3: Path = /api*                  → TG-API',
    '  │  Rule 4: Path = /web*                  → TG-Web',
    '  │  Rule 5: Query platform=mobile         → TG-Mobile',
    '  │  Rule 6: Query platform=desktop        → TG-Desktop',
    '  │  Rule 7: HTTP Method = GET             → TG-GET',
    '  │  Rule 8: HTTP Method = POST            → TG-POST',
    '  └──────────────────────────────────────────────────┘',
],'All ALB Routing Rules')

div()
step_banner(1,'Launch EC2-App1 (Host-Based Routing)')
body('**Objective:** Server for `app1.example.com`')
tbl(['Setting','Value'],[
    ['Name','EC2-App1'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['Subnet','Public-Subnet-1'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
cat <<EOF > /var/www/html/index.html
<html>
<body style='background-color:#4CAF50; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>🟢 APP 1 Server</h1>
<h2 style='color:white;'>Host: app1.example.com</h2>
</body>
</html>
EOF''')

div()
step_banner(2,'Launch EC2-App2 (Host-Based Routing)')
body('**Objective:** Server for `app2.example.com`')
tbl(['Setting','Value'],[
    ['Name','EC2-App2'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['Subnet','Public-Subnet-2'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
cat <<EOF > /var/www/html/index.html
<html>
<body style='background-color:#2196F3; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>🔵 APP 2 Server</h1>
<h2 style='color:white;'>Host: app2.example.com</h2>
</body>
</html>
EOF''')

div()
step_banner(3,'Launch EC2-API (Path-Based Routing)')
body('**Objective:** Server for `/api`')
tbl(['Setting','Value'],[
    ['Name','EC2-API'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['Subnet','Public-Subnet-1'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
mkdir -p /var/www/html/api
cat <<EOF > /var/www/html/api/index.html
<html>
<body style='background-color:#FF9800; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>🟠 API Server</h1>
<h2 style='color:white;'>Path: /api</h2>
</body>
</html>
EOF''')

div()
step_banner(4,'Launch EC2-Web (Path-Based Routing)')
body('**Objective:** Server for `/web`')
tbl(['Setting','Value'],[
    ['Name','EC2-Web'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['Subnet','Public-Subnet-2'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
mkdir -p /var/www/html/web
cat <<EOF > /var/www/html/web/index.html
<html>
<body style='background-color:#9C27B0; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>🟣 Web Server</h1>
<h2 style='color:white;'>Path: /web</h2>
</body>
</html>
EOF''')

div()
step_banner(5,'Launch EC2-Mobile (Query String Routing)')
body('**Objective:** Server for `?platform=mobile`')
tbl(['Setting','Value'],[
    ['Name','EC2-Mobile'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['Subnet','Public-Subnet-1'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
cat <<EOF > /var/www/html/index.html
<html>
<body style='background-color:#E91E63; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>📱 Mobile Server</h1>
<h2 style='color:white;'>You are on Mobile Platform!</h2>
<h3 style='color:white;'>Routed via Query String</h3>
</body>
</html>
EOF''')

div()
step_banner(6,'Launch EC2-Desktop (Query String Routing)')
body('**Objective:** Server for `?platform=desktop`')
tbl(['Setting','Value'],[
    ['Name','EC2-Desktop'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['Subnet','Public-Subnet-2'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
cat <<EOF > /var/www/html/index.html
<html>
<body style='background-color:#607D8B; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>🖥️ Desktop Server</h1>
<h2 style='color:white;'>You are on Desktop Platform!</h2>
<h3 style='color:white;'>Routed via Query String</h3>
</body>
</html>
EOF''')

div()
step_banner(7,'Create All Target Groups')
tbl(['Target Group','EC2 Target','Purpose'],[
    ['TG-App1','EC2-App1','Host-Based Routing'],
    ['TG-App2','EC2-App2','Host-Based Routing'],
    ['TG-API','EC2-API','Path-Based Routing'],
    ['TG-Web','EC2-Web','Path-Based Routing'],
    ['TG-Mobile','EC2-Mobile','Query String Routing'],
    ['TG-Desktop','EC2-Desktop','Query String Routing'],
],col_widths=[1.8,1.8,2.9])

div()
step_banner(8,'Create Application Load Balancer')
tbl(['Setting','Value'],[
    ['Name','My-ALB'],['Type','Internet-Facing'],
    ['Protocol','HTTP'],['Port','80'],['Security Group','LB-SG'],
    ['Default Action','Forward to TG-Desktop'],
],col_widths=[2.5,4.0])

div()
step_banner(9,'Configure Route 53 DNS Records')
tbl(['Record Name','Points To'],[
    ['example.com','My-ALB'],
    ['app1.example.com','My-ALB'],
    ['app2.example.com','My-ALB'],
],col_widths=[3.0,3.5])
info('All DNS records point to the SAME ALB. The ALB reads the Host Header to differentiate between them.')

div()
step_banner(10,'Configure Host-Based Routing Rules')
h4('Rule 1 — App1')
tbl(['Setting','Value'],[
    ['Host Header','app1.example.com'],['Action','Forward to TG-App1'],['Priority','1'],
],col_widths=[2.5,4.0])
h4('Rule 2 — App2')
tbl(['Setting','Value'],[
    ['Host Header','app2.example.com'],['Action','Forward to TG-App2'],['Priority','2'],
],col_widths=[2.5,4.0])

div()
step_banner(11,'Configure Path-Based Routing Rules')
h4('API Rule')
tbl(['Setting','Value'],[['Path','/api*'],['Action','Forward to TG-API'],['Priority','3']],col_widths=[2.5,4.0])
h4('Web Rule')
tbl(['Setting','Value'],[['Path','/web*'],['Action','Forward to TG-Web'],['Priority','4']],col_widths=[2.5,4.0])

div()
step_banner(12,'Configure Query String Routing Rules')
h4('Mobile Query Rule')
tbl(['Setting','Value'],[
    ['Query Key','platform'],['Query Value','mobile'],
    ['Action','Forward to TG-Mobile'],['Priority','7'],
],col_widths=[2.5,4.0])
h4('Desktop Query Rule')
tbl(['Setting','Value'],[
    ['Query Key','platform'],['Query Value','desktop'],
    ['Action','Forward to TG-Desktop'],['Priority','8'],
],col_widths=[2.5,4.0])

div()
step_banner(13,'Launch EC2-GET (HTTP Method Routing)')
body('**Objective:** Handle GET requests.')
tbl(['Setting','Value'],[
    ['Name','EC2-GET'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['Subnet','Public-Subnet-1'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
cat <<EOF > /var/www/html/index.html
<html>
<body style='background-color:#00BCD4; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>📖 GET Server - Read Operations</h1>
<h2 style='color:white;'>You sent a GET Request!</h2>
<h3 style='color:white;'>This server handles READ operations only</h3>
</body>
</html>
EOF''')

div()
step_banner(14,'Launch EC2-POST (HTTP Method Routing)')
body('**Objective:** Handle POST requests.')
tbl(['Setting','Value'],[
    ['Name','EC2-POST'],['AMI','Ubuntu'],['Instance Type','t2.micro'],
    ['Subnet','Public-Subnet-2'],['Security Group','EC2-SG'],
],col_widths=[2.5,4.0])
h4('User Data')
cod('''#!/bin/bash
apt update -y
apt install apache2 -y
systemctl start apache2
systemctl enable apache2
cat <<EOF > /var/www/html/index.html
<html>
<body style='background-color:#F44336; text-align:center;'>
<h1 style='color:white; margin-top:200px;'>✏️ POST Server - Write Operations</h1>
<h2 style='color:white;'>You sent a POST Request!</h2>
<h3 style='color:white;'>This server handles WRITE operations only</h3>
</body>
</html>
EOF''')

div()
step_banner(15,'Create Target Groups for HTTP Method Routing')
tbl(['Target Group','Purpose'],[
    ['TG-GET','GET Requests'],['TG-POST','POST Requests'],
],col_widths=[2.5,4.0])

div()
step_banner(16,'Configure HTTP Method Routing Rules')
h4('GET Rule')
tbl(['Setting','Value'],[['HTTP Method','GET'],['Action','Forward to TG-GET'],['Priority','9']],col_widths=[2.5,4.0])
h4('POST Rule')
tbl(['Setting','Value'],[['HTTP Method','POST'],['Action','Forward to TG-POST'],['Priority','10']],col_widths=[2.5,4.0])

div()
step_banner(17,'Test Host-Based Routing')
cod('http://app1.example.com  →  Expected: GREEN APP1 Page 🟢\nhttp://app2.example.com  →  Expected: BLUE APP2 Page 🔵')
arch([
    '  app1.example.com ──Host Header──▶ ALB ──Rule 1 matches──▶ TG-App1 → EC2-App1 🟢',
    '  app2.example.com ──Host Header──▶ ALB ──Rule 2 matches──▶ TG-App2 → EC2-App2 🔵',
],'Host-Based Routing Test')

div()
step_banner(18,'Test Path-Based Routing')
cod('http://example.com/api  →  Expected: ORANGE API Page 🟠\nhttp://example.com/web  →  Expected: PURPLE WEB Page 🟣')
arch([
    '  example.com/api ──▶ ALB ──/api* matches (Priority 3)──▶ TG-API → EC2-API 🟠',
    '  example.com/web ──▶ ALB ──/web* matches (Priority 4)──▶ TG-Web → EC2-Web 🟣',
],'Path-Based Routing Test')

div()
step_banner(19,'Test Query String Routing')
cod('http://example.com?platform=mobile   →  Expected: PINK Mobile Page 📱\nhttp://example.com?platform=desktop  →  Expected: GREY Desktop Page 🖥️')
arch([
    '  ?platform=mobile  ──▶ ALB ──Query matches (Priority 7)──▶ TG-Mobile  → EC2-Mobile  📱',
    '  ?platform=desktop ──▶ ALB ──Query matches (Priority 8)──▶ TG-Desktop → EC2-Desktop 🖥️',
],'Query String Routing Test')

div()
step_banner(20,'Test HTTP Method Routing')
cod('curl -X GET  http://example.com  →  Expected: GET Server Page 📖\ncurl -X POST http://example.com  →  Expected: POST Server Page ✏️')
arch([
    '  GET  Request ──▶ ALB ──Method=GET  (Priority 9)──▶  TG-GET  → EC2-GET  📖',
    '  POST Request ──▶ ALB ──Method=POST (Priority 10)──▶ TG-POST → EC2-POST ✏️',
    '',
    '  Real-world use: GET = Read servers (cache-optimised)',
    '                 POST = Write servers (DB-write optimised)',
],'HTTP Method Routing Test')

div()
h3('Routing Methods Summary')
tbl(['Routing Type','Rule Condition','Priority','Target Group','Status'],[
    ['Host-Based','app1.example.com','1','TG-App1','✅ Completed'],
    ['Host-Based','app2.example.com','2','TG-App2','✅ Completed'],
    ['Path-Based','/api*','3','TG-API','✅ Completed'],
    ['Path-Based','/web*','4','TG-Web','✅ Completed'],
    ['Query String','platform=mobile','7','TG-Mobile','✅ Completed'],
    ['Query String','platform=desktop','8','TG-Desktop','✅ Completed'],
    ['HTTP Method','GET','9','TG-GET','✅ Completed'],
    ['HTTP Method','POST','10','TG-POST','✅ Completed'],
],col_widths=[1.7,1.9,0.9,1.4,1.1])

cod('''✅ Host-Based Routing
✅ Path-Based Routing
✅ Query String Routing
✅ HTTP Method Routing
✅ Route 53 + ALB Integration
✅ Multi-Target Group Architecture''')

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — NLB BANKING TASK (existing)
# ═══════════════════════════════════════════════════════════════════════════════
div()
section_banner('AWS NLB Implementation — Real-World Banking TCP Gateway','Network Load Balancer · Layer 4 · TCP Port 9000')

div()
h2('Scenario')
body('This demo demonstrates a real-world TCP-based banking gateway architecture where:')
bul('Application Load Balancer (ALB) cannot be used')
bul('Network Load Balancer (NLB) is required')
bul('Raw TCP traffic is load balanced')
bul('Health checks and failover are demonstrated')
body('**ALB only supports:** HTTP, HTTPS, WebSocket, gRPC')
body('**Our application uses:** Raw TCP traffic on port 9000')
info('Therefore: ALB cannot support this architecture. NLB works at Layer 4 (Transport Layer) and supports TCP, UDP, TLS.')

arch([
    '  ATM Client (EC2: 10.0.2.132)',
    '        │  TCP : 9000',
    '        ▼',
    '  Network Load Balancer (banking-nlb)',
    '  DNS: banking-nlb-xxxx.elb.ap-south-1.amazonaws.com',
    '        │',
    '  ──────────────────────────────────',
    '  │                                │',
    '  ▼                                ▼',
    'banking-core-1              banking-core-2',
    '10.0.1.101                  10.0.2.102',
    'TCP Banking Server          TCP Banking Server',
    'AZ: ap-south-1a             AZ: ap-south-1b',
],'NLB Banking Gateway Architecture')

div()
h2('Step-by-Step Implementation')
step_banner(1,'Create VPC','16213E')
tbl(['Setting','Value'],[['VPC Name','Banking-VPC'],['CIDR','10.0.0.0/16']],col_widths=[2.5,4.0])

step_banner(2,'Create Public Subnets','16213E')
tbl(['Subnet Name','CIDR','Availability Zone'],[
    ['Public-Subnet-A','10.0.1.0/24','ap-south-1a'],
    ['Public-Subnet-B','10.0.2.0/24','ap-south-1b'],
],col_widths=[2.0,2.0,2.5])

step_banner(3,'Create Internet Gateway','16213E')
tbl(['Setting','Value'],[['Internet Gateway Name','Banking-IGW'],['Attach to','Banking-VPC']],col_widths=[2.5,4.0])

step_banner(4,'Create Route Table','16213E')
tbl(['Destination','Target'],[['10.0.0.0/16','local'],['0.0.0.0/0','Banking-IGW']],col_widths=[2.5,4.0])
body('**Associate Route Table With:** Public-Subnet-A, Public-Subnet-B')

step_banner(5,'Create Security Groups','16213E')
h4('banking-backend-sg — Attached to banking-core-1 and banking-core-2')
tbl(['Type','Port','Protocol','Source'],[
    ['SSH','22','TCP','Your Public IP'],
    ['Custom TCP','9000','TCP','0.0.0.0/0'],
],col_widths=[1.5,1.0,1.2,2.8])
h4('atm-client-sg — Attached to atm-client')
tbl(['Type','Port','Protocol','Source'],[['SSH','22','TCP','Your Public IP']],col_widths=[1.5,1.0,1.2,2.8])
note('Traditionally: NLB does NOT use Security Groups. Traffic filtering is handled at Backend EC2 Security Groups.')

step_banner(6,'Create EC2 Instances','16213E')
tbl(['Instance','Purpose','Subnet'],[
    ['banking-core-1','Backend TCP Banking Server','Public-Subnet-A'],
    ['banking-core-2','Backend TCP Banking Server','Public-Subnet-B'],
    ['atm-client','Simulated ATM Client','Any Public Subnet'],
],col_widths=[1.8,2.5,2.2])
body('**AMI:** Ubuntu 22.04  ·  **Instance Type:** t2.micro  ·  **Public IP:** Enabled')
h4('User Data — banking-core-1')
cod('''#!/bin/bash
apt update -y
apt install netcat-openbsd -y
cat <<EOF > /home/ubuntu/server1.sh
#!/bin/bash
while true
do
    echo "Connected to Banking Server 1 at AZ A - Transaction Gateway" | nc -l -p 9000 -N
done
EOF
chmod +x /home/ubuntu/server1.sh
cat <<EOF > /etc/systemd/system/banking-server1.service
[Unit]
Description=Banking TCP Server 1
After=network.target
[Service]
ExecStart=/home/ubuntu/server1.sh
Restart=always
User=root
[Install]
WantedBy=multi-user.target
EOF
systemctl daemon-reload
systemctl enable banking-server1.service
systemctl start banking-server1.service''')
h4('User Data — banking-core-2')
cod('''#!/bin/bash
apt update -y
apt install netcat-openbsd -y
cat <<EOF > /home/ubuntu/server2.sh
#!/bin/bash
while true
do
    echo "Connected to Banking Server 2 at AZ B - Transaction Gateway" | nc -l -p 9000 -N
done
EOF
chmod +x /home/ubuntu/server2.sh
cat <<EOF > /etc/systemd/system/banking-server2.service
[Unit]
Description=Banking TCP Server 2
After=network.target
[Service]
ExecStart=/home/ubuntu/server2.sh
Restart=always
User=root
[Install]
WantedBy=multi-user.target
EOF
systemctl daemon-reload
systemctl enable banking-server2.service
systemctl start banking-server2.service''')
h4('User Data — atm-client')
cod('#!/bin/bash\napt update -y\napt install netcat-openbsd -y')

step_banner(7,'Verify Backend Servers','16213E')
body('SSH into: **atm-client**')
cod('''nc <BANKING-CORE-1-PUBLIC-IP> 9000
Expected: Connected to Banking Server 1 at Availability Zone A - Transaction Gateway

nc <BANKING-CORE-2-PUBLIC-IP> 9000
Expected: Connected to Banking Server 2 at Availability Zone B - Transaction Gateway''')

step_banner(8,'Create Target Group','16213E')
tbl(['Setting','Value'],[
    ['Target Group Name','banking-tg'],['Target Type','Instances'],
    ['Protocol','TCP'],['Port','9000'],['Health Check Protocol','TCP'],
    ['Healthy Threshold','3'],['Unhealthy Threshold','3'],
    ['Interval','10s'],['Timeout','6s'],
    ['Register Targets','banking-core-1, banking-core-2'],
],col_widths=[2.5,4.0])

step_banner(9,'Create Network Load Balancer','16213E')
tbl(['Setting','Value'],[
    ['Name','banking-nlb'],['Scheme','Internet-facing'],
    ['Listener Protocol','TCP'],['Listener Port','9000'],
    ['Attach Subnets','Public-Subnet-A (ap-south-1a), Public-Subnet-B (ap-south-1b)'],
    ['Attach Target Group','banking-tg'],
],col_widths=[2.5,4.0])
body('Wait until status becomes: **Active**')

step_banner(10,'Test NLB','16213E')
cod('''nc <NLB-DNS> 9000
Expected: Connected to Banking Server 1 ...

nc <NLB-DNS> 9000
Expected: Connected to Banking Server 2 ...''')
h4('Live Load Balancing Demo')
cod('''while true
do
  nc <NLB-DNS> 9000
  sleep 1
done

Expected Output:
Server 1
Server 2
Server 1
Server 2

This demonstrates: TCP load balancing in real time.''')

step_banner(11,'Failover Demo','16213E')
body('Stop: **banking-core-1**  ·  Wait: 30–40 seconds')
cod('''nc <NLB-DNS> 9000

Only Server 2 responds.
This demonstrates: automatic health checks and failover.''')

step_banner(12,'Source IP Preservation Demo','16213E')
body('SSH into backend server **banking-core-1**, install and start tcpdump:')
cod('''sudo apt install tcpdump -y
sudo tcpdump -i any port 9000''')
body('On atm-client, check private IP and connect:')
cod('''hostname -I
# Example: 10.0.2.132

nc <NLB-DNS> 9000''')
body('Observe tcpdump output:')
cod('''IP 10.0.2.132.54872 > 10.0.1.25.9000

This proves: NLB preserves the original client IP.''')

div()
h2('Important NLB Concepts')
h3('NLB Works at Layer 4')
body('Supports: TCP, UDP, TLS')
h3('NLB Preserves Source IP')
body('Backend servers can see: **Real Client IP**')
body('Useful in: Banking, Fraud Detection, Security Logging')
h3('NLB is Built For:')
bul('Banking Systems'); bul('Gaming Servers'); bul('Trading Platforms')
bul('IoT Applications'); bul('TCP-Based Applications'); bul('Ultra-Low Latency Workloads')

div()
p=doc.add_paragraph();sp(p,b=10,a=0);pshd(p,'0F3460')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
ir(p,'"Application Load Balancer understands applications.\nNetwork Load Balancer understands network traffic."',
   size=13,bold=True,italic=True,color=WHITE)
p=doc.add_paragraph();sp(p,b=0,a=0);pshd(p,'16213E')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
ir(p,'Documentation prepared for AWS Load Balancer Case Study Module',size=9,color=RGBColor(0xAA,0xCC,0xFF))

doc.save('lb-case-study.docx')
print('Done → lb-case-study.docx')
