import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    colors = {1:(15,52,96), 2:(22,33,62), 3:(15,52,96), 4:(83,52,131)}
    sizes  = {1:20, 2:15, 3:12, 4:11}
    set_font(run, size=sizes[level], bold=True, color=colors[level])
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_paragraph(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, color=(15,52,96))
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=10, color=(192,57,43))
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_font(run)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A1A2E')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=9, color=(220,220,220))
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=10, bold=True, color=(255,255,255))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0F3460')
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx+1]
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, size=10, bold=True, color=(15,52,96))
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    set_font(run, name="Courier New", size=9, color=(192,57,43))
                else:
                    run = p.add_run(part)
                    set_font(run, size=10)
            if r_idx % 2 == 1:
                tc = tr.cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F4FF')
                tcPr.append(shd)
    doc.add_paragraph()

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level*0.2)
    p.paragraph_format.space_after = Pt(3)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, color=(15,52,96))
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=10, color=(192,57,43))
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_font(run)

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8F4FD')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, italic=True, color=(15,52,96))

def banner(text, sub='', fill='0F3460', txt_color=(255,255,255)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f'  {text}  ')
    set_font(r, size=22, bold=True, color=txt_color)
    if sub:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr2 = p2._p.get_or_add_pPr()
        shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'), 'clear')
        shd2.set(qn('w:color'), 'auto')
        shd2.set(qn('w:fill'), 'E94560')
        pPr2.append(shd2)
        p2.paragraph_format.space_after = Pt(16)
        r2 = p2.add_run(f'  {sub}  ')
        set_font(r2, size=10, color=(255,255,255))

def section_banner(text, fill='16213E'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(f'  {text}')
    set_font(r, size=14, bold=True, color=(255,255,255))

def add_image(path, width_in=5.5, caption=''):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        p.add_run().add_picture(path, width=Inches(width_in))
    except Exception as e:
        add_paragraph(f'[Diagram: {caption}]')
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        cr = cp.add_run(caption)
        set_font(cr, size=8.5, italic=True, color=(80,80,120))

def add_step(num, title, body=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run(f'Step {num}: ')
    set_font(r1, bold=True, color=(83,52,131))
    r2 = p.add_run(title)
    set_font(r2, bold=True, color=(15,52,96))
    if body:
        add_paragraph(body, indent=0.3)

# ─────────────────────────────────────────────────────────────────────────────
# COVER
# ─────────────────────────────────────────────────────────────────────────────
banner(
    'AWS Certificate Manager (ACM)',
    'TLS/SSL Fundamentals · Certificate Providers · ACM Deep Dive · Hands-On Task'
)

# ─────────────────────────────────────────────────────────────────────────────
# PART 1 — SSL / TLS FUNDAMENTALS
# ─────────────────────────────────────────────────────────────────────────────
section_banner('PART 1 — SSL / TLS Fundamentals', '0F3460')

add_heading('What is SSL and TLS?', 1)
add_paragraph(
    'Every time you see the padlock icon in a browser, **TLS (Transport Layer Security)** is '
    'silently protecting your connection. It ensures three things:'
)
add_bullet('**Confidentiality** — data is encrypted so no one in the middle can read it')
add_bullet('**Integrity** — data cannot be modified in transit without detection')
add_bullet('**Authentication** — you are talking to the real server, not an impersonator')

add_paragraph(
    '**SSL (Secure Sockets Layer)** was the original protocol created by Netscape in 1994. '
    'All versions of SSL are now deprecated and considered insecure. '
    '**TLS** is its successor, first standardized in 1999, and is what the entire internet '
    'uses today. Despite this, the industry still commonly uses "SSL" as a colloquial term '
    'for what is technically TLS.'
)

add_heading('SSL/TLS Version History', 2)
add_table(
    ['Version', 'Year', 'Status', 'Notes'],
    [
        ['SSL 1.0', '1994', 'Never released', 'Internal Netscape draft — serious security flaws found'],
        ['SSL 2.0', '1995', 'Deprecated (RFC 6176)', 'First public release — broken, MITM vulnerable'],
        ['SSL 3.0', '1996', 'Deprecated (RFC 7568)', 'POODLE attack — do not use'],
        ['TLS 1.0', '1999', 'Deprecated (2021)', 'Upgrade of SSL 3.0 — BEAST attack vulnerability'],
        ['TLS 1.1', '2006', 'Deprecated (2021)', 'Minor fixes — still not recommended'],
        ['TLS 1.2', '2008', 'Widely supported', 'Still secure with correct configuration'],
        ['**TLS 1.3**', '**2018**', '**Current standard**', '**Faster, simpler, more secure — use this**'],
    ]
)
add_blockquote(
    'Modern AWS services (ALB, CloudFront, ACM) support TLS 1.2 and TLS 1.3. '
    'TLS 1.3 reduces handshake latency and removes insecure cipher suites entirely.'
)

add_heading('What is an X.509 Certificate?', 2)
add_paragraph(
    'An **X.509 certificate** is the standard format for SSL/TLS certificates. '
    'It is a digital document that binds a public key to an identity (domain name). '
    'It is signed by a Certificate Authority (CA) to prove its authenticity.'
)
add_paragraph('A certificate contains:')
add_table(
    ['Field', 'Example', 'Purpose'],
    [
        ['**Subject**', 'CN=saistabakers.com, O=Sais Bakers', 'Who the cert belongs to'],
        ['**SAN (Subject Alt Name)**', 'DNS:saistabakers.com, DNS:www.saistabakers.com', 'All domains covered by this cert'],
        ['**Issuer**', 'Amazon RSA 2048 M01', 'Which CA signed this cert'],
        ['**Public Key**', 'RSA 2048-bit or EC P-256', 'Used in TLS handshake key exchange'],
        ['**Validity Period**', 'Not Before / Not After dates', 'When the cert is valid'],
        ['**Serial Number**', 'Unique hex string', 'CA\'s unique identifier for this cert'],
        ['**Signature**', 'CA\'s cryptographic signature', 'Proves the CA verified and signed it'],
        ['**Key Usage**', 'digitalSignature, keyEncipherment', 'What the cert can be used for'],
    ]
)

add_heading('Certificate Chain of Trust', 2)
add_paragraph(
    'Certificates work in a **chain of trust**. Your browser comes pre-installed with '
    'a set of **Root CA certificates** (from companies like DigiCert, Comodo, Amazon, '
    'Let\'s Encrypt\'s ISRG). When your browser receives a certificate, it walks the '
    'chain from your domain\'s certificate up to a trusted root.'
)
add_image('diagram_cert_chain.png', width_in=5.5,
          caption='X.509 Certificate Chain — Root CA → Intermediate CA → Leaf Certificate → Browser Trust')

add_paragraph(
    'If any link in the chain is missing, self-signed, or expired, the browser shows '
    '**NET::ERR_CERT_AUTHORITY_INVALID** and warns users not to proceed.'
)

# ─────────────────────────────────────────────────────────────────────────────
# SSL/TLS CERTIFICATE PROVIDERS
# ─────────────────────────────────────────────────────────────────────────────
add_heading('Certificate Providers and Certificate Types', 1)

add_heading('Certificate Validation Levels', 2)
add_table(
    ['Type', 'Validation', 'Browser Shows', 'Best For', 'Typical Cost'],
    [
        ['**DV — Domain Validated**', 'Prove domain ownership only (DNS or HTTP file)', 'Padlock', 'Blogs, personal sites, APIs', 'Free (Let\'s Encrypt) to $10/yr'],
        ['**OV — Organization Validated**', 'Domain + company identity verified by CA', 'Padlock + org in cert details', 'Business websites', '$50–$200/yr'],
        ['**EV — Extended Validation**', 'Rigorous legal entity verification', 'Padlock (was green bar, now removed)', 'Banks, e-commerce, legal', '$100–$500/yr'],
        ['**Wildcard**', 'DV or OV, covers *.example.com', 'Padlock', 'Subdomains at scale', '$70–$300/yr'],
        ['**Multi-Domain (SAN)**', 'Multiple domains on one cert', 'Padlock', 'Multiple products/brands', 'Varies'],
    ]
)

add_heading('Major Certificate Providers', 2)
add_table(
    ['Provider', 'Cost', 'Validity', 'Automation', 'Notes'],
    [
        ["**Let's Encrypt**", 'Free', '90 days', 'ACME protocol / certbot', 'Largest CA by volume. Non-profit. No EV/OV.'],
        ['**AWS ACM (Public)**', 'Free', '13 months (auto-renew)', 'Fully automated', 'Only for AWS services (ALB, CloudFront, API GW)'],
        ['**DigiCert**', '$175–$500+/yr', '1–2 years', 'Partial', 'Premium CA. EV, OV, DV. Enterprise SLAs.'],
        ['**Comodo / Sectigo**', '$70–$300/yr', '1–2 years', 'Partial', 'Wide browser support. Wildcard & multi-domain.'],
        ['**GlobalSign**', '$150–$400+/yr', '1–2 years', 'API available', 'Enterprise, IoT, DevOps focus.'],
        ['**ZeroSSL**', 'Free / Paid', '90 days (free)', 'ACME support', "Let's Encrypt alternative. REST API."],
        ['**Google Trust Services**', 'Free (via GCP)', '90 days', 'ACME', 'For Google Cloud users.'],
        ["**Self-Signed**", 'Free', 'Any', 'Manual', 'Development only — not trusted by browsers.'],
    ]
)

add_heading("How to Get an SSL Certificate — Let's Encrypt (Without Cloud)", 2)
add_paragraph(
    "Let's Encrypt is the most widely used free CA. It uses the **ACME protocol** "
    "(Automated Certificate Management Environment, RFC 8555) to automate domain validation "
    "and certificate issuance. The tool **certbot** handles this automatically."
)
add_code_block(
    '# Ubuntu/Debian — Install certbot\n'
    'sudo apt update\n'
    'sudo apt install certbot python3-certbot-apache -y\n\n'
    '# Obtain certificate and auto-configure Apache\n'
    'sudo certbot --apache -d example.com -d www.example.com\n\n'
    '# Certbot will:\n'
    '#  1. Create a temp file at /.well-known/acme-challenge/<token>\n'
    '#  2. Let\'s Encrypt CA fetches that URL to verify domain ownership\n'
    '#  3. CA issues certificate, certbot installs it in Apache\n'
    '#  4. Adds cron job / systemd timer for auto-renewal every 60 days\n\n'
    '# DNS validation (when no port 80 access)\n'
    'sudo certbot certonly --manual --preferred-challenges dns -d example.com\n'
    '# Certbot gives you a TXT record to add to your DNS provider\n\n'
    '# Test renewal\n'
    'sudo certbot renew --dry-run\n\n'
    '# Certificate files stored at:\n'
    '/etc/letsencrypt/live/example.com/fullchain.pem   # cert + chain\n'
    '/etc/letsencrypt/live/example.com/privkey.pem     # private key'
)

add_heading('How to Get an SSL Certificate — Manual / Paid CA', 2)
for num, (title, body) in enumerate([
    ('Generate Private Key + CSR on your server',
     'The private key never leaves your server. The CSR (Certificate Signing Request) '
     'contains your public key and domain info and is sent to the CA.'),
    ('Submit CSR to CA (DigiCert, Comodo, etc.)', ''),
    ('Complete Domain Validation (DV) or Organization Validation (OV/EV)',
     'For DV: add a CNAME/TXT DNS record or upload a file to your web server. '
     'For OV/EV: the CA calls your company and verifies legal documents.'),
    ('Download Certificate Files from CA',
     'You receive: certificate.crt (your cert), ca-bundle.crt (intermediate chain), '
     'and your existing private key.'),
    ('Install Certificate on Server', ''),
], 1):
    add_step(num, title, body)

add_code_block(
    '# Step 1 — Generate private key and CSR\n'
    'openssl genrsa -out private.key 2048\n'
    'openssl req -new -key private.key -out domain.csr \\\n'
    '  -subj "/C=IN/ST=Maharashtra/L=Mumbai/O=Sais Bakers/CN=saistabakers.com"\n\n'
    '# View CSR content\n'
    'openssl req -text -noout -in domain.csr\n\n'
    '# Step 5 — Install on Nginx\n'
    'ssl_certificate     /etc/ssl/certs/certificate.crt;\n'
    'ssl_certificate_key /etc/ssl/private/private.key;\n'
    'ssl_trusted_certificate /etc/ssl/certs/ca-bundle.crt;\n\n'
    '# Install on Apache\n'
    'SSLCertificateFile    /etc/ssl/certs/certificate.crt\n'
    'SSLCertificateKeyFile /etc/ssl/private/private.key\n'
    'SSLCACertificateFile  /etc/ssl/certs/ca-bundle.crt'
)

# ─────────────────────────────────────────────────────────────────────────────
# TLS HANDSHAKE
# ─────────────────────────────────────────────────────────────────────────────
add_heading('TLS Handshake — How a Secure Connection is Established', 1)

add_paragraph(
    'Before any application data is exchanged, the **TLS handshake** establishes '
    'a shared secret between client and server, authenticates the server\'s identity '
    'via its certificate, and negotiates the encryption algorithms to use.'
)

add_image('diagram_tls_handshake.png', width_in=6.2,
          caption='TLS 1.3 Handshake — Full Message Flow with Key Derivation')

add_heading('TLS 1.3 Handshake Step by Step', 2)
add_table(
    ['Step', 'Message', 'Who Sends', 'What It Contains'],
    [
        ['1', '**ClientHello**', 'Client → Server',
         'Supported TLS versions, cipher suites, client random nonce, key_share (EC public key)'],
        ['2', '**ServerHello**', 'Server → Client',
         'Chosen cipher suite, server random nonce, server key_share, session ID'],
        ['3', '**EncryptedExtensions**', 'Server → Client',
         'Additional negotiated parameters (ALPN, SNI confirmation)'],
        ['4', '**Certificate**', 'Server → Client',
         'Server\'s X.509 cert chain (leaf + intermediate CA)'],
        ['5', '**CertificateVerify**', 'Server → Client',
         'Digital signature over handshake transcript proving server holds the private key'],
        ['6', '**Finished**', 'Server → Client',
         'HMAC over entire handshake — proves integrity of negotiation'],
        ['7', '**Finished**', 'Client → Server',
         'Client HMAC — confirms both sides derived the same keys'],
        ['8', '**Application Data**', 'Both (encrypted)',
         'All payload encrypted with AEAD cipher (AES-256-GCM or ChaCha20-Poly1305)'],
    ]
)

add_heading('Key Derivation (HKDF)', 3)
add_paragraph(
    'TLS 1.3 uses **HKDF (HMAC-based Key Derivation Function)** to derive multiple keys '
    'from the shared secret produced by the (EC)DHE key exchange:'
)
add_code_block(
    'Inputs:\n'
    '  Client Random  +  Server Random  +  (EC)DHE shared secret\n\n'
    'HKDF derives:\n'
    '  handshake_traffic_secret  → encrypts the handshake itself\n'
    '  client_application_traffic_secret  → client write key + IV\n'
    '  server_application_traffic_secret  → server write key + IV\n'
    '  exporter_master_secret    → for applications needing keying material\n\n'
    'Session keys are EPHEMERAL — new keys for every session\n'
    'Perfect Forward Secrecy: past sessions cannot be decrypted even if private key is later compromised'
)

add_heading('How Payload is Encrypted', 2)
add_paragraph(
    'TLS uses a combination of **asymmetric cryptography** (for key exchange) and '
    '**symmetric cryptography** (for bulk data encryption):'
)
add_table(
    ['Phase', 'Algorithm Type', 'Algorithm Used', 'Purpose'],
    [
        ['Key Exchange', 'Asymmetric', 'ECDHE (P-256, X25519)', 'Securely establish shared secret without transmitting it'],
        ['Authentication', 'Asymmetric', 'RSA-2048 / ECDSA P-256', 'Server proves identity using private key signature'],
        ['Key Derivation', 'Symmetric', 'HKDF-SHA256 / HKDF-SHA384', 'Derive session keys from shared secret'],
        ['Data Encryption', 'Symmetric', 'AES-256-GCM / ChaCha20-Poly1305', 'Encrypt all application data (fast, AEAD)'],
        ['Integrity', 'Built into AEAD', 'Poly1305 / GCM tag', 'Detect any tampering with encrypted data'],
    ]
)
add_blockquote(
    'AEAD = Authenticated Encryption with Associated Data. It simultaneously encrypts '
    'the payload AND produces an authentication tag. If any byte is altered in transit, '
    'decryption fails and the connection is dropped immediately.'
)

# ─────────────────────────────────────────────────────────────────────────────
# TLS TERMINATION AT LOAD BALANCER
# ─────────────────────────────────────────────────────────────────────────────
add_heading('TLS Termination at the Load Balancer', 1)

add_paragraph(
    '**TLS Termination** means the Load Balancer decrypts incoming HTTPS traffic '
    'and forwards plain HTTP to the backend servers. This is the standard pattern '
    'in AWS and most cloud architectures.'
)

add_heading('Why Terminate TLS at the Load Balancer?', 2)
add_table(
    ['Reason', 'Explanation'],
    [
        ['**Certificate in one place**', 'Single cert on the ALB — backends need no cert management'],
        ['**Offload CPU cost**', 'TLS crypto is CPU-intensive. ALB handles it, freeing EC2 for app logic'],
        ['**Inspect HTTP headers**', 'ALB can route by path, host, header — only possible after decryption'],
        ['**WAF integration**', 'AWS WAF inspects decrypted HTTP traffic for SQL injection, XSS, etc.'],
        ['**Centralized renewal**', 'ACM auto-renews one cert instead of managing certs on every EC2'],
        ['**Scaling**', 'Add/remove EC2s without redistributing certificates'],
    ]
)

add_code_block(
    'HTTPS Traffic Flow with TLS Termination:\n\n'
    'Browser\n'
    '  │\n'
    '  │  HTTPS :443  (TLS encrypted)\n'
    '  ▼\n'
    'ALB  ←── ACM Certificate attached here\n'
    '  │  TLS Handshake + Decryption happens at ALB\n'
    '  │\n'
    '  │  HTTP :80  (plain text — travels over private VPC network only)\n'
    '  ▼\n'
    'EC2 Instances  (no cert needed, receives plain HTTP)\n\n'
    'Optional — End-to-End Encryption (E2E TLS):\n'
    'Browser → HTTPS :443 → ALB → HTTPS :443 → EC2\n'
    '(ALB re-encrypts to backend — used for strict compliance requirements)\n\n'
    'Optional — mTLS (Mutual TLS):\n'
    'Client presents its own cert → ALB validates client cert → Backend\n'
    '(Used in zero-trust, B2B APIs, service mesh)'
)

add_heading('Traditional SSL vs Cloud SSL — Flow Comparison', 2)
add_image('diagram_traditional_ssl.png', width_in=6.0,
          caption='Traditional SSL/TLS — Self-Managed Certificate on Your Server (Without Cloud)')

add_paragraph(
    'In the traditional model you generate a private key on your server, create a CSR, '
    'submit it to a CA, download the signed certificate, install it manually, configure '
    'your web server, and remember to renew it before it expires. The private key never '
    'leaves your server — but managing this across many servers becomes complex and error-prone.'
)

# ─────────────────────────────────────────────────────────────────────────────
# PART 2 — AWS ACM
# ─────────────────────────────────────────────────────────────────────────────
section_banner('PART 2 — AWS Certificate Manager (ACM)', '0F3460')

add_heading('What is AWS Certificate Manager?', 1)

add_paragraph(
    '**AWS Certificate Manager (ACM)** is a fully managed service that provisions, manages, '
    'and automatically renews SSL/TLS certificates for use with AWS services. '
    'ACM eliminates the manual work of certificate management — you never handle private keys, '
    'never manually renew, and certificates are free for use with supported AWS services.'
)

add_image('diagram_aws_acm_ssl.png', width_in=6.2,
          caption='AWS ACM — Cloud SSL Certificate Flow: Issuance · Validation · Attachment · Traffic')

add_heading('ACM Key Features', 2)
add_table(
    ['Feature', 'Details'],
    [
        ['**Free for AWS services**', 'No charge for ACM public certificates. Only pay for ACM Private CA.'],
        ['**Automatic renewal**', 'ACM renews certs at least 60 days before expiry — completely hands-free'],
        ['**Private key security**', 'Private key stored in AWS HSMs (Hardware Security Modules) — you never see it'],
        ['**Two validation methods**', 'DNS validation (recommended) or Email validation'],
        ['**Wildcard support**', 'Single cert for *.example.com covers unlimited subdomains'],
        ['**Multi-region**', 'CloudFront requires cert in us-east-1. ALB requires cert in same region.'],
        ['**Transparency logging**', 'All issued certs logged to public Certificate Transparency logs'],
        ['**ACM Private CA**', 'Create your own internal CA for private services, mTLS, IoT'],
    ]
)

add_heading('ACM vs Self-Managed Certificate', 2)
add_table(
    ['Aspect', 'ACM (Cloud)', 'Self-Managed (Let\'s Encrypt / DigiCert)'],
    [
        ['Cost', 'Free (for AWS resources)', 'Free (LE) to $500+/yr (DigiCert EV)'],
        ['Renewal', 'Fully automatic', 'Manual or cron job (can fail)'],
        ['Private Key Access', 'Never exposed — stored in AWS HSM', 'You have full access to key'],
        ['Installation', 'Click to attach to ALB/CloudFront', 'Manual server config (Nginx/Apache)'],
        ['Works on non-AWS servers', 'No — AWS services only', 'Yes — any server'],
        ['Wildcard', 'Yes', 'Yes'],
        ['EV/OV certificates', 'No — DV only', 'Yes (paid CAs)'],
        ['Private CA', 'ACM Private CA (paid)', 'Self-hosted (OpenSSL, CFSSL)'],
        ['Validity', '13 months (auto-renewed)', '90 days (LE) or 1-2 years (paid)'],
    ]
)

add_heading('ACM Core Components', 1)

add_heading('1. Public Certificate', 3)
add_paragraph(
    'A free, domain-validated certificate issued by **Amazon Trust Services** (Amazon\'s own CA). '
    'Can only be attached to: **ALB, NLB, CloudFront, API Gateway, Elastic Beanstalk, '
    'AppSync, and other integrated AWS services**. '
    'Cannot be exported or installed on an EC2 directly.'
)

add_heading('2. Certificate Validation', 3)
add_paragraph('ACM must verify you own the domain before issuing a certificate:')
add_table(
    ['Method', 'How It Works', 'Best For'],
    [
        ['**DNS Validation (Recommended)**',
         'ACM gives you a CNAME record to add to your DNS. '
         'ACM polls the record periodically. Auto-renews as long as the CNAME stays in DNS.',
         'Route 53 users (one-click add), programmatic workflows'],
        ['**Email Validation**',
         'ACM emails the domain\'s WHOIS contacts (admin@, webmaster@, etc.). '
         'Someone must click the approval link.',
         'Domains where you cannot modify DNS'],
    ]
)

add_heading('3. Amazon Trust Services (ATS)', 3)
add_paragraph(
    '**Amazon Trust Services** is Amazon\'s own Certificate Authority, launched in 2016. '
    'It operates the root CAs and intermediate CAs that sign all ACM-issued certificates.'
)
add_table(
    ['Root CA', 'Algorithm', 'Validity', 'Trusted Since'],
    [
        ['Amazon Root CA 1', 'RSA 2048', '2038', '2016 — pre-installed in all major browsers/OS'],
        ['Amazon Root CA 2', 'RSA 4096', '2040', '2016'],
        ['Amazon Root CA 3', 'EC P-256', '2040', '2016'],
        ['Amazon Root CA 4', 'EC P-384', '2040', '2016'],
        ['Starfield Services Root CA G2', 'RSA 2048', '2037', 'Cross-certified by older Starfield root for legacy devices'],
    ]
)
add_paragraph(
    'ACM certificates are signed by an **Intermediate CA** (e.g. Amazon RSA 2048 M01, M02) '
    'which is in turn signed by Amazon Root CA 1. Browsers trust Amazon Root CA 1 by default, '
    'so ACM certs are trusted everywhere.'
)

add_heading('4. Automatic Renewal', 3)
add_paragraph(
    'ACM begins the renewal process at least **60 days before expiry** for DNS-validated certs. '
    'As long as the validation CNAME record remains in your DNS, renewal is fully automatic '
    'and zero-downtime — the new certificate is silently swapped in.'
)
add_blockquote(
    'For email-validated certs, someone must click the renewal email within 72 hours. '
    'This is why DNS validation is strongly recommended for production workloads.'
)

add_heading('5. ACM Private CA', 3)
add_paragraph(
    '**ACM Private CA** (previously AWS Private CA) lets you create and operate your own '
    'private Certificate Authority for issuing certificates to internal services:'
)
add_bullet('Private microservices communicating over mTLS inside a VPC')
add_bullet('IoT devices that need unique certificates')
add_bullet('Internal tools that need HTTPS without public CA validation')
add_bullet('Zero-trust architectures requiring client certificates')
add_paragraph('Pricing: $400/month per CA + $0.75 per certificate issued.')

add_heading('ACM Supported AWS Services', 2)
add_table(
    ['Service', 'Notes'],
    [
        ['**Application Load Balancer (ALB)**', 'Most common — HTTPS listener attaches ACM cert'],
        ['**Network Load Balancer (NLB)**', 'TLS listener on port 443'],
        ['**CloudFront**', 'Certificate must be in us-east-1 (N. Virginia) region'],
        ['**API Gateway**', 'Custom domain name with ACM cert'],
        ['**Elastic Beanstalk**', 'HTTPS on load balancer'],
        ['**AppSync**', 'Custom domain with ACM cert'],
        ['**Cognito**', 'Custom domain for hosted UI'],
        ['**Global Accelerator**', 'TLS termination at edge'],
    ]
)
add_blockquote(
    'ACM certificates CANNOT be directly installed on EC2 instances. '
    'If you need HTTPS directly on EC2 (without a load balancer), use Let\'s Encrypt or a paid CA.'
)

add_heading('ACM Certificate Lifecycle Flow', 2)
add_code_block(
    'You request cert in ACM console\n'
    '  │\n'
    '  │ Enter domain: saistabakers.com + *.saistabakers.com\n'
    '  │ Choose: DNS validation\n'
    '  ▼\n'
    'ACM generates CNAME record\n'
    '  _abc123xyz.saistabakers.com → _xyz789abc.acm-validations.aws\n'
    '  │\n'
    '  │ You add this CNAME to Route 53 (or any DNS provider)\n'
    '  ▼\n'
    'ACM polls your DNS every few minutes\n'
    '  │\n'
    '  │ Sees CNAME record → Domain ownership confirmed\n'
    '  ▼\n'
    'Amazon Trust Services signs certificate\n'
    '  │\n'
    '  │ Status changes: Pending Validation → Issued\n'
    '  ▼\n'
    'You attach cert to ALB HTTPS listener\n'
    '  │\n'
    '  │ Select cert from ACM in ALB listener config\n'
    '  ▼\n'
    'HTTPS live on your domain\n'
    '  │\n'
    '  │ Auto-renewed at 60 days before expiry\n'
    '  ▼\n'
    'Renewal: ACM polls same CNAME → Re-issues → Zero downtime swap'
)

# ─────────────────────────────────────────────────────────────────────────────
# PART 3 — TASK
# ─────────────────────────────────────────────────────────────────────────────
section_banner('PART 3 — Hands-On Task', '533483')

add_heading('Task — Enable HTTPS on saistabakers.com Using ACM + Route 53', 1)

add_heading('Task Overview', 2)
add_paragraph(
    'In this task, you already own the domain **saistabakers.com** purchased from **Hostinger**. '
    'The goal is to serve a secure HTTPS Apache website using AWS Route 53 for DNS and '
    'AWS ACM for the SSL certificate — without managing any certificate files manually.'
)

add_table(
    ['Component', 'Detail'],
    [
        ['Domain', 'saistabakers.com (purchased on Hostinger)'],
        ['DNS Provider', 'Amazon Route 53 (Hosted Zone created, NS updated at Hostinger)'],
        ['SSL Certificate', 'AWS ACM — DNS validated, free, auto-renewing'],
        ['Validation', 'CNAME record added in Route 53'],
        ['Traffic', 'Simple Routing — A record → EC2 Public IP'],
        ['Web Server', 'Apache2 on EC2 serving HTTPS (via ACM on ALB or directly)'],
        ['Subdomain', 'www.saistabakers.com → same EC2 (additional CNAME in Route 53)'],
    ]
)

add_heading('Task Architecture Diagram', 2)
add_image('diagram_task_architecture.png', width_in=6.4,
          caption='saistabakers.com — Hostinger → Route 53 → ACM DNS Validation → HTTPS Apache on EC2')

add_heading('Complete Step-by-Step Guide', 2)

# Step 1
add_step(1, 'Create Route 53 Hosted Zone for saistabakers.com')
add_bullet('Open AWS Console → Route 53 → Hosted Zones → **Create Hosted Zone**')
add_bullet('Domain name: `saistabakers.com`')
add_bullet('Type: **Public Hosted Zone**')
add_bullet('Click Create')
add_paragraph('Route 53 auto-creates NS and SOA records. Note the 4 NS records:', indent=0.3)
add_code_block(
    'ns-xxx.awsdns-xx.com\n'
    'ns-yyy.awsdns-yy.net\n'
    'ns-zzz.awsdns-zz.org\n'
    'ns-www.awsdns-ww.co.uk'
)

# Step 2
add_step(2, 'Update Nameservers at Hostinger')
add_bullet('Log in to Hostinger → Domains → saistabakers.com → **DNS / Nameservers**')
add_bullet('Change nameservers from Hostinger default to **Custom Nameservers**')
add_bullet('Enter all 4 Route 53 NS records')
add_bullet('Save — DNS propagation takes a few minutes to 48 hours')
add_blockquote(
    'Once Hostinger nameservers are updated, all DNS for saistabakers.com is '
    'controlled by Route 53. Any records you add in Route 53 are now live on the internet.'
)

# Step 3
add_step(3, 'Launch EC2 Instance and Install Apache2')
add_table(
    ['Setting', 'Value'],
    [
        ['AMI', 'Ubuntu 22.04 LTS (or Amazon Linux 2023)'],
        ['Instance Type', 't2.micro (Free Tier)'],
        ['Key Pair', 'Create or select existing'],
        ['Public IP', 'Enable auto-assign'],
        ['Security Group', 'Allow SSH :22 (your IP), HTTP :80 (anywhere), HTTPS :443 (anywhere)'],
    ]
)
add_paragraph('Connect and install Apache2:', indent=0.3)
add_code_block(
    'ssh -i your-key.pem ubuntu@<EC2-PUBLIC-IP>\n\n'
    'sudo apt update -y\n'
    'sudo apt install apache2 -y\n'
    'sudo systemctl start apache2\n'
    'sudo systemctl enable apache2\n\n'
    '# Create a simple webpage\n'
    "sudo bash -c 'cat > /var/www/html/index.html << EOF\n"
    '<!DOCTYPE html>\n'
    '<html>\n'
    '<head><title>Sais Bakers</title></head>\n'
    '<body>\n'
    '<h1>Welcome to Sais Bakers!</h1>\n'
    '<p>Secured with AWS ACM + Route 53</p>\n'
    '</body>\n'
    '</html>\n'
    "EOF'\n\n"
    '# Verify Apache is running\n'
    'sudo systemctl status apache2\n'
    'curl http://localhost'
)
add_paragraph('Note your EC2 public IP — you\'ll need it for the DNS A record.', indent=0.3)

# Step 4
add_step(4, 'Request SSL Certificate in AWS ACM')
add_bullet('Open AWS Console → **Certificate Manager (ACM)**')
add_bullet('Click **Request a certificate** → **Request a public certificate**')
add_bullet('Add domain names:')
add_code_block('saistabakers.com\n*.saistabakers.com')
add_bullet('Validation method: **DNS validation** (Recommended)', 1)
add_bullet('Key algorithm: RSA 2048', 1)
add_bullet('Click **Request**')
add_paragraph(
    'The certificate status shows **Pending validation**. ACM is waiting for you to prove '
    'domain ownership by adding a CNAME record to your DNS.', indent=0.3
)

# Step 5
add_step(5, 'Add ACM CNAME Validation Record to Route 53')
add_paragraph(
    'In the ACM certificate details page, you will see a CNAME record that ACM generated:', indent=0.3
)
add_code_block(
    'Record Name:  _abc123xyz.saistabakers.com\n'
    'Record Type:  CNAME\n'
    'Record Value: _xyz789abc.acm-validations.aws'
)
add_paragraph('Add this to Route 53:', indent=0.3)
add_bullet('Go to Route 53 → Hosted Zones → saistabakers.com')
add_bullet('Click **Create Record**')
add_bullet('Record name: `_abc123xyz` (just the part before .saistabakers.com)')
add_bullet('Record type: **CNAME**')
add_bullet('Value: `_xyz789abc.acm-validations.aws`')
add_bullet('TTL: 300')
add_bullet('Click **Create records**')
add_paragraph(
    'Shortcut: If your domain is already in Route 53, ACM offers a '
    '**"Create records in Route 53"** button that adds the CNAME automatically in one click.', indent=0.3
)
add_paragraph('Wait 2–5 minutes. ACM detects the CNAME and the status changes to **Issued**.', indent=0.3)

# Step 6
add_step(6, 'Add Subdomain CNAME for www.saistabakers.com')
add_paragraph(
    'Add a second DNS record in Route 53 to handle the `www` subdomain:', indent=0.3
)
add_code_block(
    'Record Name:  www\n'
    'Record Type:  CNAME\n'
    'Record Value: saistabakers.com\n'
    'TTL:          300'
)
add_paragraph(
    'This makes `www.saistabakers.com` resolve to the same destination as `saistabakers.com`. '
    'The ACM wildcard cert (`*.saistabakers.com`) already covers the www subdomain.', indent=0.3
)

# Step 7
add_step(7, 'Create A Record with Simple Routing — Point Domain to EC2')
add_bullet('Route 53 → Hosted Zones → saistabakers.com → **Create Record**')
add_table(
    ['Field', 'Value'],
    [
        ['Record Name', '(leave blank — for root domain saistabakers.com)'],
        ['Record Type', 'A — Routes traffic to IPv4 address'],
        ['Routing Policy', 'Simple routing'],
        ['Value', 'EC2 Public IP (e.g. 54.12.34.56)'],
        ['TTL', '300'],
    ]
)
add_paragraph('Click **Create records**.', indent=0.3)
add_paragraph(
    'Your Route 53 hosted zone now has these records:', indent=0.3
)
add_code_block(
    'saistabakers.com              A      54.12.34.56           (→ EC2)\n'
    'www.saistabakers.com          CNAME  saistabakers.com      (→ root)\n'
    '_abc123.saistabakers.com      CNAME  _xyz789.acm-valid...  (ACM validation)\n'
    'saistabakers.com              NS     ns-xxx.awsdns-xx.com  (auto-created)\n'
    'saistabakers.com              SOA    (auto-created)'
)

# Step 8
add_step(8, 'Attach ACM Certificate to ALB (for full HTTPS end-to-end)')
add_paragraph(
    'Since ACM certs cannot be directly installed on EC2, you need an **Application Load Balancer** '
    'to terminate TLS and forward HTTP to EC2. If you want HTTPS directly to EC2 without an ALB, '
    'use Let\'s Encrypt certbot instead (see Step 8B below).', indent=0.3
)
add_paragraph('ALB path (Recommended for production):', indent=0.3)
add_bullet('Create ALB in same region as your EC2')
add_bullet('Add **HTTPS :443 listener** → select your ACM certificate')
add_bullet('Add **HTTP :80 listener** → redirect to HTTPS (optional)')
add_bullet('Create Target Group → add EC2 → attach to ALB')
add_bullet('Update Route 53 A record → change value to ALB DNS name (use **Alias record**)')
add_code_block(
    '# Route 53 after ALB setup:\n'
    'saistabakers.com   A (Alias)  →  my-alb-1234.us-east-1.elb.amazonaws.com\n\n'
    '# Security Group: EC2\n'
    'Allow HTTP :80 from ALB Security Group (not 0.0.0.0/0)\n\n'
    '# Security Group: ALB\n'
    'Allow HTTPS :443 from 0.0.0.0/0\n'
    'Allow HTTP  :80  from 0.0.0.0/0'
)

add_paragraph('Alternative — Step 8B: certbot directly on EC2 (no ALB needed):', indent=0.3)
add_code_block(
    '# Install certbot on EC2\n'
    'sudo apt install certbot python3-certbot-apache -y\n\n'
    '# Get certificate (requires port 80 open and DNS already pointing to EC2)\n'
    'sudo certbot --apache -d saistabakers.com -d www.saistabakers.com\n\n'
    '# certbot auto-configures Apache with HTTPS and sets up auto-renewal\n'
    'sudo systemctl reload apache2\n\n'
    '# Verify HTTPS\n'
    'curl -I https://saistabakers.com'
)

# Step 9
add_step(9, 'Validate the Complete Setup')
add_code_block(
    '# Check DNS propagation\n'
    'nslookup saistabakers.com\n'
    'dig saistabakers.com A\n'
    'dig www.saistabakers.com CNAME\n\n'
    '# Check certificate details\n'
    'openssl s_client -connect saistabakers.com:443 -servername saistabakers.com\n\n'
    '# Check certificate info\n'
    'curl -vI https://saistabakers.com 2>&1 | grep -E "SSL|certificate|issuer"\n\n'
    '# Expected output:\n'
    '* Server certificate:\n'
    '*  subject: CN=saistabakers.com\n'
    '*  start date: ...\n'
    '*  expire date: ...\n'
    '*  issuer: C=US, O=Amazon, CN=Amazon RSA 2048 M01\n'
    '*  SSL certificate verify ok.\n\n'
    '# Check browser:\n'
    'https://saistabakers.com   →  Padlock visible\n'
    'https://www.saistabakers.com  →  Same page'
)

add_heading('DNS Records Summary', 2)
add_table(
    ['Record Name', 'Type', 'Value', 'Purpose'],
    [
        ['saistabakers.com', 'A', '54.12.34.56 (or ALB Alias)', 'Simple routing → EC2 / ALB'],
        ['www.saistabakers.com', 'CNAME', 'saistabakers.com', 'www subdomain → root domain'],
        ['_abc123.saistabakers.com', 'CNAME', '_xyz789.acm-validations.aws', 'ACM DNS validation'],
        ['saistabakers.com', 'NS', '4× awsdns NS records', 'Route 53 nameservers (auto)'],
        ['saistabakers.com', 'SOA', 'ns-xxx.awsdns-xx.com', 'Zone authority record (auto)'],
    ]
)

add_heading('Common Errors and Fixes', 2)
add_table(
    ['Error', 'Cause', 'Fix'],
    [
        ['ACM cert stuck in Pending Validation',
         'CNAME not yet added or DNS not propagated',
         'Verify CNAME in Route 53. Wait 5 min. Check with: dig _abc123.saistabakers.com CNAME'],
        ['Site shows insecure / padlock missing',
         'HTTP still served or cert not attached',
         'Check ALB HTTPS listener. Redirect HTTP→HTTPS. Verify ACM cert status is Issued.'],
        ['nslookup returns Hostinger IPs',
         'Hostinger nameservers not yet updated or propagation pending',
         'Check Hostinger panel. NS records must match Route 53 NS. Wait up to 48 hours.'],
        ['ERR_CERT_NAME_INVALID',
         'Cert domain does not match the URL',
         'Cert was issued for saistabakers.com but accessed via IP or different domain.'],
        ['Apache serves HTTP not HTTPS on EC2 directly',
         'No cert on EC2 — ACM only works with ALB',
         'Use ALB with ACM cert, or use certbot on EC2 directly.'],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# KEY TAKEAWAYS
# ─────────────────────────────────────────────────────────────────────────────
add_heading('Key Takeaways', 1)

add_bullet('**TLS** replaced SSL — TLS 1.3 is the current standard. Never use SSL 2.0/3.0 or TLS 1.0/1.1.')
add_bullet('**X.509 certificates** bind a domain to a public key and are signed by a trusted CA.')
add_bullet('**Certificate chain of trust**: Root CA → Intermediate CA → Leaf certificate.')
add_bullet('**TLS Handshake** authenticates the server, negotiates cipher, and derives session keys — all before data flows.')
add_bullet('**AEAD encryption** (AES-256-GCM, ChaCha20-Poly1305) ensures both confidentiality and integrity of payload.')
add_bullet('**TLS Termination at ALB** offloads certificate management and enables HTTP-level routing and WAF inspection.')
add_bullet('**ACM** is free, fully managed, and auto-renews. Best choice for all AWS-hosted applications.')
add_bullet('**Amazon Trust Services** is Amazon\'s own CA — its root certs are trusted by all major browsers.')
add_bullet('**DNS Validation** is preferred over Email Validation — it enables fully automatic renewal.')
add_bullet('**ACM certs cannot be installed on EC2 directly** — use ALB, or use certbot/Let\'s Encrypt for direct EC2 HTTPS.')
add_bullet('**Wildcard cert** (`*.saistabakers.com`) covers all subdomains with one certificate.')

# Closing banner
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), '0F3460')
pPr.append(shd)
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run(
    '  HTTPS is not optional — ACM makes it effortless. '
    'Secure every endpoint, automate every renewal.  '
)
set_font(r, size=11, italic=True, color=(210, 225, 255))

doc.save('acm-case-study.docx')
print('Done → acm-case-study.docx')
