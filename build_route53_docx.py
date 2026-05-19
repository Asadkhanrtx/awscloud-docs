import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        set_font(run, size=20, bold=True, color=(15, 52, 96))
    elif level == 2:
        set_font(run, size=15, bold=True, color=(22, 33, 62))
    elif level == 3:
        set_font(run, size=12, bold=True, color=(15, 52, 96))
    elif level == 4:
        set_font(run, size=11, bold=True, color=(83, 52, 131))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_paragraph(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, color=(15, 52, 96))
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=10, color=(192, 57, 43))
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_font(run)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A1A2E')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=9, color=(220, 220, 220))
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0F3460')
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, size=10, bold=True, color=(15, 52, 96))
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    set_font(run, name="Courier New", size=9, color=(192, 57, 43))
                else:
                    run = p.add_run(part)
                    set_font(run, size=10)
            if r_idx % 2 == 1:
                tc = tr.cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F4FF')
                tcPr.append(shd)
    doc.add_paragraph()

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    p.paragraph_format.space_after = Pt(3)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, color=(15, 52, 96))
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=10, color=(192, 57, 43))
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_font(run)

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8F4FD')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, italic=True, color=(15, 52, 96))

def add_diagram_label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '533483')
    pPr.append(shd)
    run = p.add_run(f'  {text}  ')
    set_font(run, size=9, bold=True, color=(255, 255, 255))

def add_image(path, width_inches=5.5, caption=None):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(8)
            cr = cp.add_run(caption)
            set_font(cr, size=9, italic=True, color=(100, 100, 100))
    except Exception as e:
        add_paragraph(f'[Architecture Diagram: {caption or path}]')

def add_numbered(text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    first = True
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run((f'{num}. ' if first else '') + part[2:-2])
            set_font(run, bold=True, color=(15, 52, 96))
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run((f'{num}. ' if first else '') + part[1:-1])
            set_font(run, name="Courier New", size=10, color=(192, 57, 43))
        else:
            run = p.add_run((f'{num}. ' if first else '') + part)
            set_font(run)
        first = False

# ─────────────────────────────────────────────────────────────
# COVER BANNER
# ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), '0F3460')
pPr.append(shd)
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(6)
r = p.add_run('  AWS ROUTE 53 & DNS  ')
set_font(r, size=26, bold=True, color=(255, 255, 255))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr2 = p2._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:val'), 'clear')
shd2.set(qn('w:color'), 'auto')
shd2.set(qn('w:fill'), 'E94560')
pPr2.append(shd2)
p2.paragraph_format.space_after = Pt(20)
r2 = p2.add_run('  Complete Case Study — DNS Fundamentals · Route 53 · Routing Policies · Practical Tasks  ')
set_font(r2, size=11, bold=False, color=(255, 255, 255))

# ─────────────────────────────────────────────────────────────
# SECTION 0 — HOW THE INTERNET FINDS YOUR WEBSITE
# ─────────────────────────────────────────────────────────────
add_heading('How the Internet Finds Your Website', 1)
add_paragraph(
    'Every time you type a website address into a browser, a complex chain of lookups happens '
    'invisibly within milliseconds. Understanding this chain — DNS — is the foundation for '
    'everything Route 53 does.'
)

add_heading('What Happens When a Client Hits a Domain', 2)
add_paragraph(
    'When you type `www.example.com` and press Enter, your device does not know the IP address '
    'of that server. It needs to ask. The Domain Name System (DNS) is the global distributed '
    'database that maps human-readable names to machine-readable IP addresses. Here is the full '
    'journey step by step:'
)

add_diagram_label('ARCHITECTURE — Full DNS Resolution Flow')
add_code_block(
    'Client Browser\n'
    '      │\n'
    '      │ 1. Check Browser Cache\n'
    '      │    (is example.com already cached? TTL still valid?)\n'
    '      │\n'
    '      │ 2. Check OS / hosts file\n'
    '      │    (/etc/hosts, Windows hosts file)\n'
    '      │\n'
    '      ▼\n'
    'Recursive Resolver  ←──── (ISP DNS / 8.8.8.8 / 1.1.1.1)\n'
    '      │\n'
    '      │ 3. Check Resolver Cache\n'
    '      │    (has it resolved this recently?)\n'
    '      │\n'
    '      │ 4. Ask Root DNS Server\n'
    '      │    "Who handles .com?"\n'
    '      ▼\n'
    'Root DNS Server  (13 root server clusters worldwide)\n'
    '      │\n'
    '      │ "Go ask the .com TLD server"\n'
    '      │ Returns address of .com TLD nameserver\n'
    '      ▼\n'
    'TLD DNS Server  (.com / .org / .net / .io etc.)\n'
    '      │\n'
    '      │ "Go ask the Authoritative NS for example.com"\n'
    '      │ Returns the Authoritative Name Server records\n'
    '      ▼\n'
    'Authoritative Name Server  (Route 53 / Cloudflare / etc.)\n'
    '      │\n'
    '      │ "www.example.com → 54.12.34.56"\n'
    '      │ Returns the actual A Record / IP address\n'
    '      ▼\n'
    'Recursive Resolver  (caches result for TTL duration)\n'
    '      │\n'
    '      ▼\n'
    'Client Browser\n'
    '      │\n'
    '      │ 5. TCP/TLS connection established to 54.12.34.56\n'
    '      │ 6. HTTP request sent\n'
    '      ▼\n'
    'Web Server  →  Response  →  Page loads\n'
    '\n'
    'Result also cached in: Browser Cache + OS Cache + Resolver Cache (based on TTL)'
)

add_heading('What is a Nameserver?', 3)
add_paragraph(
    'A **nameserver** is a server that stores DNS records and answers DNS queries. When you '
    'register a domain and create a hosted zone, Route 53 assigns you 4 nameservers. These '
    'are the servers that the rest of the internet will ask when looking up your domain.'
)
add_paragraph(
    'Nameservers form a hierarchy: **Root Nameservers** → **TLD Nameservers** → '
    '**Authoritative Nameservers**. Each layer delegates responsibility to the next layer '
    'until the actual record is found.'
)
add_diagram_label('ARCHITECTURE — Nameserver Hierarchy')
add_code_block(
    'ROOT SERVERS  (.)  ← 13 globally distributed clusters, managed by ICANN\n'
    '     │\n'
    '     ├── .com  TLD Servers  (managed by Verisign)\n'
    '     │      │\n'
    '     │      └── example.com  Authoritative NS  (managed by Route 53)\n'
    '     │               │\n'
    '     │               ├── www.example.com  →  54.12.34.56   (A Record)\n'
    '     │               ├── api.example.com  →  54.12.34.99   (A Record)\n'
    '     │               └── mail.example.com →  mail.google   (MX Record)\n'
    '     │\n'
    '     ├── .org  TLD Servers\n'
    '     ├── .net  TLD Servers\n'
    '     └── .io   TLD Servers'
)

add_heading('DNS Delegation — How Your Domain Points to Route 53', 3)
add_paragraph(
    'When you buy a domain from GoDaddy or Namecheap, their nameservers are used by default. '
    'To use Route 53 as your DNS provider, you delegate authority by updating the nameservers '
    'at your registrar to point to Route 53\'s nameservers.'
)
add_diagram_label('ARCHITECTURE — DNS Delegation Flow')
add_code_block(
    'Domain Registrar (GoDaddy / Namecheap)\n'
    '     │\n'
    '     │  You update NS records here to point to Route 53\n'
    '     ▼\n'
    'Route 53 Nameservers\n'
    '     ns-123.awsdns-45.com\n'
    '     ns-678.awsdns-12.net\n'
    '     ns-901.awsdns-34.org\n'
    '     ns-234.awsdns-56.co.uk\n'
    '     │\n'
    '     │  Route 53 now answers all DNS queries for your domain\n'
    '     ▼\n'
    'DNS Queries from the internet → Route 53 Hosted Zone → DNS Records → IP Address'
)

# ─────────────────────────────────────────────────────────────
# SECTION 1 — INTRODUCTION TO DNS
# ─────────────────────────────────────────────────────────────
add_heading('Introduction to DNS', 1)

add_heading('What is DNS?', 2)
add_paragraph(
    'DNS (Domain Name System) translates human-friendly domain names into machine-readable IP addresses.'
)
add_paragraph('Example:')
add_code_block('www.google.com  →  172.217.18.36')
add_paragraph(
    'DNS acts like the **internet\'s phone book**, allowing users to access websites using names '
    'instead of remembering IP addresses.'
)

add_heading('Why DNS is Important', 2)
add_paragraph('Without DNS, users would need to remember IP addresses for every website.')
add_paragraph('DNS provides:')
add_bullet('Easy-to-remember website names')
add_bullet('Stable access even if server IPs change')
add_bullet('Scalability for millions of websites')
add_bullet('Load distribution through multiple A records (Round Robin DNS)')
add_bullet('Security and spam protection through TXT / SPF / DKIM records')

add_heading('DNS Hierarchy', 2)
add_paragraph('DNS uses a hierarchical naming structure:')
add_code_block(
    '.\n'
    '└── .com\n'
    '    └── example.com\n'
    '        ├── www.example.com\n'
    '        └── api.example.com'
)

add_heading('How DNS Resolution Works', 2)

add_heading('DNS Resolution Steps', 3)
add_paragraph('When you enter `www.example.com` in a browser:')

steps = [
    ('Browser Cache Check', 'Browser checks if the domain IP is already cached locally.'),
    ('OS Cache Check', 'Operating system checks its DNS cache (and the local hosts file).'),
    ('Query Recursive Resolver', 'Device sends the request to a recursive resolver (ISP, Google DNS `8.8.8.8`, Cloudflare `1.1.1.1`).'),
    ('Root DNS Server', 'Resolver asks the Root Server for `.com` information. Root Server points to the `.com` TLD server.'),
    ('TLD DNS Server', 'Resolver asks the TLD server for `example.com`. TLD server returns the Authoritative Name Server.'),
    ('Authoritative Name Server', 'Returns the actual IP address of `www.example.com`.'),
    ('Browser Connects to Server', 'Browser receives the IP and loads the website.'),
    ('Caching', 'The result is cached for future requests based on TTL.'),
]
for num, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run_bold = p.add_run(f'{num}. {title} — ')
    set_font(run_bold, bold=True, color=(15, 52, 96))
    run_desc = p.add_run(desc.replace('`', ''))
    set_font(run_desc)

add_heading('Key DNS Components', 3)
add_table(
    ['Component', 'Purpose'],
    [
        ['**Recursive Resolver**', 'Performs the full DNS lookup on behalf of the client. Caches results.'],
        ['**Root Server**', 'Directs requests to TLD servers. 13 root server clusters worldwide.'],
        ['**TLD Server**', 'Handles domains like .com, .org, .net. Knows which NS is authoritative.'],
        ['**Authoritative Server**', 'Stores the actual DNS records for a domain. Final answer source.'],
        ['**DNS Cache**', 'Speeds up future lookups by storing recent results based on TTL.'],
    ]
)

add_heading('Important DNS Terms', 3)
add_table(
    ['Term', 'Description'],
    [
        ['**Domain Name**', 'Human-readable website name (e.g. example.com)'],
        ['**IP Address**', 'Numerical server address (e.g. 54.12.34.56 for IPv4)'],
        ['**TTL**', 'Time To Live — how long DNS records stay cached before re-querying'],
        ['**FQDN**', 'Fully Qualified Domain Name — complete domain name (www.example.com.)'],
        ['**TLD**', 'Top-Level Domain — the last segment of a domain name (.com, .org, .io)'],
        ['**Zone File**', 'Text file containing all DNS records for a hosted zone'],
        ['**Propagation**', 'The time it takes for DNS changes to spread across global servers'],
        ['**Delegation**', 'Assigning DNS authority for a subdomain to a different nameserver'],
    ]
)

# ─────────────────────────────────────────────────────────────
# SECTION 2 — DNS SERVICE PROVIDERS
# ─────────────────────────────────────────────────────────────
add_heading('DNS Service Providers', 1)

add_heading('Popular DNS Providers', 2)
add_table(
    ['Provider', 'Type', 'Features'],
    [
        ['**GoDaddy**', 'Registrar + DNS Hosting', 'Beginner-friendly, popular domain provider'],
        ['**Namecheap**', 'Registrar + DNS Hosting', 'Affordable pricing, free Whois privacy'],
        ['**Cloudflare**', 'DNS Hosting + CDN', 'Fast DNS, DDoS protection, free tier, proxy mode'],
        ['**AWS Route 53**', 'Registrar + DNS Hosting', 'AWS integration, advanced routing policies, health checks'],
        ['**Google Cloud DNS**', 'DNS Hosting', 'Low latency, global anycast network'],
    ]
)

add_heading('Domain Registrar vs DNS Hosting', 2)

add_heading('Domain Registrar', 3)
add_paragraph('Lets you purchase and own a domain name and maintains domain ownership records.')
add_paragraph('Examples:')
add_bullet('GoDaddy')
add_bullet('Namecheap')
add_bullet('Route 53')

add_heading('DNS Hosting Provider', 3)
add_paragraph('Hosts DNS records for your domain and responds to DNS queries from the internet.')
add_paragraph('Examples:')
add_bullet('Cloudflare')
add_bullet('Route 53')
add_bullet('Google Cloud DNS')

add_heading('Simple Analogy', 3)
add_bullet('**Domain Registrar** → Registers your house address with the city')
add_bullet('**DNS Hosting Provider** → Delivers mail to your house once the address is known')
add_paragraph(
    'You can buy a domain from **GoDaddy** and use **Route 53** or **Cloudflare** for DNS hosting. '
    'The registrar and the DNS host do not have to be the same company — you simply point your '
    'registrar\'s NS records to whichever DNS host you prefer.'
)

# ─────────────────────────────────────────────────────────────
# SECTION 3 — AWS ROUTE 53
# ─────────────────────────────────────────────────────────────
add_heading('AWS Route 53', 1)

add_heading('Architecture — Route 53 in a Production AWS Setup', 2)
add_image('/home/user/awscloud-docs/architecture.png', width_inches=5.8,
          caption='AWS Route 53 → ALB → Auto Scaling EC2 Instances across Multiple Availability Zones')

add_heading('What is Route 53?', 2)
add_paragraph(
    'Amazon Route 53 is a **highly available, scalable, fully managed, and authoritative DNS service** '
    'provided by AWS.'
)
add_blockquote(
    'Authoritative DNS means you can directly manage and update DNS records for your domain. '
    'Route 53\'s answer is the final, definitive answer — it is not a cached or forwarded response.'
)
add_paragraph('Route 53 also acts as:')
add_bullet('A **Domain Registrar** — purchase and manage domains directly through AWS')
add_bullet('A **DNS Routing Service** — routes traffic using intelligent policies')
add_bullet('A **Health Checking Service** — monitors endpoints and reroutes on failure')

add_heading('Main Features', 2)
add_table(
    ['Feature', 'Description'],
    [
        ['**Domain Registration**', 'Purchase and manage domains through AWS. Integrates instantly with hosted zones.'],
        ['**DNS Routing**', 'Routes traffic to resources like EC2, S3, ALB, CloudFront, API Gateway'],
        ['**Health Checks**', 'Monitors endpoints and routes traffic only to healthy resources'],
        ['**High Availability**', 'AWS provides a 100% availability SLA — the only AWS service with a full uptime guarantee'],
        ['**Traffic Flow**', 'Visual editor for complex multi-policy routing configurations'],
        ['**DNSSEC**', 'DNS Security Extensions — protects against spoofing and cache poisoning attacks'],
        ['**Private DNS**', 'Internal DNS resolution within VPCs without exposing records to the public internet'],
    ]
)

add_heading('Why is it Called Route 53?', 2)
add_paragraph('The name has two meanings:')
add_bullet('**Route** → Routes internet/DNS traffic to the correct destination')
add_bullet('**53** → DNS uses **port 53** for queries and responses')
add_paragraph('Port reference:')
add_code_block(
    'HTTP  → Port 80\n'
    'HTTPS → Port 443\n'
    'DNS   → Port 53   ← Route 53 is named after this'
)

# ─────────────────────────────────────────────────────────────
# SECTION 4 — ROUTE 53 CORE COMPONENTS
# ─────────────────────────────────────────────────────────────
add_heading('Route 53 Core Components', 1)

add_heading('Hosted Zones', 2)
add_paragraph(
    'A **Hosted Zone** is a container for DNS records of a domain. It is the equivalent of a DNS '
    'zone file — it holds all the records that define how traffic is routed for that domain and '
    'its subdomains.'
)
add_paragraph('Example hosted zones:')
add_bullet('`example.com`')
add_bullet('`api.example.com` (subdomain delegation)')
add_paragraph('Route 53 automatically creates two records in every new hosted zone:')
add_bullet('**NS Record** → Route 53 name servers assigned to your zone')
add_bullet('**SOA Record** → Administrative metadata about the zone')

add_heading('Public Hosted Zone', 3)
add_paragraph(
    'Used for routing traffic from the **public internet**. DNS records in a public hosted zone '
    'are visible to and resolvable by anyone on the internet.'
)
add_diagram_label('ARCHITECTURE — Public Hosted Zone')
add_code_block(
    'Internet User\n'
    '      │\n'
    '      ▼\n'
    'Public DNS Query  →  Route 53 Public Hosted Zone (example.com)\n'
    '      │\n'
    '      ├── www.example.com  →  ALB DNS name  (Alias Record)\n'
    '      ├── api.example.com  →  EC2 IP  (A Record)\n'
    '      └── mail.example.com →  mail server  (MX Record)'
)

add_heading('Private Hosted Zone', 3)
add_paragraph(
    'Used for routing traffic **inside VPCs only**. Records in a private hosted zone are only '
    'resolvable from EC2 instances and services running inside the associated VPC. '
    'They are completely invisible from the public internet.'
)
add_diagram_label('ARCHITECTURE — Private Hosted Zone')
add_code_block(
    'VPC  (10.0.0.0/16)\n'
    '      │\n'
    '      │  Internal DNS query from EC2 App Server\n'
    '      ▼\n'
    'Route 53 Private Hosted Zone (internal.example.com)\n'
    '      │\n'
    '      ├── db.internal.example.com  →  10.0.1.50  (RDS Private IP)\n'
    '      ├── cache.internal.example.com → 10.0.1.60  (ElastiCache)\n'
    '      └── app.internal.example.com  → 10.0.2.10  (Internal ALB)\n'
    '\n'
    'NOTE: These records are NOT visible from the public internet'
)
add_blockquote('Public zones work on the internet. Private zones work only inside AWS VPCs.')

add_heading('Name Servers (NS)', 2)
add_paragraph(
    'Name Servers host DNS records and answer DNS queries. Route 53 assigns 4 name servers to '
    'each hosted zone for redundancy and global availability.'
)
add_code_block(
    'ns-123.awsdns-45.com\n'
    'ns-678.awsdns-12.net\n'
    'ns-901.awsdns-34.org\n'
    'ns-234.awsdns-56.co.uk'
)
add_paragraph(
    'AWS uses **anycast routing** for its nameservers, meaning queries are automatically answered '
    'by the geographically nearest Route 53 infrastructure for lowest latency.'
)

add_heading('DNS Delegation', 2)
add_paragraph('To use Route 53 DNS for a domain purchased elsewhere:')

for num, step in enumerate([
    'Buy domain from registrar (GoDaddy, Namecheap)',
    'Create Hosted Zone in Route 53',
    'Copy the 4 Route 53 NS records',
    'Update nameservers at your registrar to the Route 53 NS values',
], 1):
    add_numbered(step, num)

add_code_block(
    'GoDaddy Domain\n'
    '      │\n'
    '      ▼\n'
    'Update Nameservers → Route 53 NS Records\n'
    '      │\n'
    '      ▼\n'
    'Route 53 Name Servers  (all DNS queries now handled by Route 53)\n'
    '      │\n'
    '      ▼\n'
    'Route 53 Hosted Zone  →  Returns correct record  →  Client gets IP'
)

add_heading('SOA Record', 2)
add_paragraph(
    'SOA = **Start of Authority**. Contains administrative details about the hosted zone. '
    'The SOA record specifies:'
)
add_bullet('**Primary Name Server** — the main authoritative NS for the zone')
add_bullet('**Administrator Contact** — the email of the zone administrator')
add_bullet('**Serial Number** — the version number of the zone (increments on each change)')
add_bullet('**Refresh Interval** — how often secondary NS servers check for zone updates')
add_bullet('**Retry Interval** — how long secondary NS waits before retrying a failed refresh')
add_bullet('**Expire Duration** — how long secondary NS will serve stale data if primary is unreachable')
add_bullet('**Default TTL** — the default Time To Live used for DNS caching')
add_blockquote('The SOA record is usually managed automatically by Route 53.')

add_heading('TTL (Time To Live)', 2)
add_paragraph(
    'TTL defines how long DNS records stay cached in resolvers and browsers before they must '
    're-query the authoritative nameserver.'
)
add_table(
    ['TTL Type', 'Effect', 'When to Use'],
    [
        ['**High TTL** (e.g. 86400 = 24 hr)', 'Less DNS traffic, slower propagation of changes', 'Stable production records'],
        ['**Low TTL** (e.g. 60 sec)', 'Faster updates, more DNS queries, higher cost', 'Before migrations, during failover testing'],
        ['**Zero TTL**', 'No caching — always queries authoritative NS', 'Rapid change scenarios (not recommended long-term)'],
    ]
)
add_blockquote('Alias records do not allow manual TTL configuration — AWS manages TTL automatically.')

# ─────────────────────────────────────────────────────────────
# SECTION 5 — DNS RECORD TYPES
# ─────────────────────────────────────────────────────────────
add_heading('DNS Record Types', 1)

add_code_block(
    'A Record       → Domain    → IPv4 Address\n'
    'AAAA Record    → Domain    → IPv6 Address\n'
    'CNAME          → Domain    → Another Domain\n'
    'MX Record      → Domain    → Mail Server\n'
    'TXT Record     → Text Metadata / Domain Verification / Email Security\n'
    'NS Record      → Authoritative Name Servers\n'
    'PTR Record     → IP Address → Domain  (Reverse DNS)\n'
    'SRV Record     → Service Host + Port\n'
    'SOA Record     → Zone administrative information\n'
    'Alias Record   → Domain    → AWS Resource  (Route 53 specific)'
)

add_heading('A Record', 2)
add_paragraph('Maps a domain name to an **IPv4 address**. The most common DNS record type.')
add_code_block('example.com  →  54.12.34.56')
add_paragraph(
    'You can have multiple A records for the same domain pointing to different IPs. '
    'DNS resolvers will return all IPs and clients typically use round-robin to pick one.'
)

add_heading('AAAA Record', 2)
add_paragraph('Maps a domain name to an **IPv6 address**.')
add_code_block('example.com  →  2001:db8::1')
add_paragraph(
    'As IPv6 adoption grows, AAAA records are increasingly important. Many modern applications '
    'configure both A and AAAA records for dual-stack support.'
)

add_heading('CNAME Record', 2)
add_paragraph('Maps one domain name to **another domain name** (not directly to an IP).')
add_code_block('www.example.com  →  example.com')
add_table(
    ['Feature', 'Details'],
    [
        ['**Limitation**', 'Cannot be used for the root/apex domain (example.com) — only subdomains'],
        ['**Resolution**', 'Resolver follows the chain until an A/AAAA record is found'],
        ['**Common Use**', 'Subdomains like www.example.com, blog.example.com, shop.example.com'],
        ['**Chaining**', 'CNAME can point to another CNAME, but avoid deep chains (latency)'],
    ]
)

add_heading('MX Record', 2)
add_paragraph('Defines **mail servers** for a domain. MX records have a priority value — lower number = higher priority.')
add_code_block('example.com  →  Priority 1  →  aspmx.l.google.com\nexample.com  →  Priority 5  →  alt1.aspmx.l.google.com')
add_blockquote(
    'When someone sends an email to user@example.com, the MX record tells the internet '
    'which mail server should receive that email.'
)

add_heading('TXT Record', 2)
add_paragraph('Stores **text-based information** for a domain. Used for:')
add_bullet('**Domain Verification** — proves domain ownership to Google Workspace, Microsoft 365, etc.')
add_bullet('**SPF (Sender Policy Framework)** — specifies which mail servers can send on behalf of your domain')
add_bullet('**DKIM (DomainKeys Identified Mail)** — cryptographic email signature verification')
add_bullet('**DMARC** — policy for handling emails that fail SPF or DKIM checks')
add_code_block(
    '# SPF Record — allow Google mail servers\n'
    'v=spf1 include:_spf.google.com ~all\n\n'
    '# Domain verification example\n'
    'google-site-verification=abc123def456'
)

add_heading('NS Record', 2)
add_paragraph(
    'Defines the **authoritative name servers** for a domain or subdomain. '
    'Automatically created by Route 53 when you create a hosted zone. '
    'NS records can also be used for subdomain delegation — pointing a subdomain to a '
    'completely different set of nameservers.'
)

add_heading('PTR Record', 2)
add_paragraph(
    'Used for **reverse DNS lookup** — mapping an IP address back to its associated domain name. '
    'Commonly used in email server verification (to prove a mail server\'s IP belongs to the claimed domain), '
    'logging, and network troubleshooting.'
)
add_code_block('54.12.34.56  →  example.com')
add_paragraph(
    'PTR records live in special zones under `.in-addr.arpa` (IPv4) or `.ip6.arpa` (IPv6). '
    'For AWS resources, you need to contact AWS support or use your ISP to set PTR records.'
)

add_heading('SRV Record', 2)
add_paragraph(
    'Defines the **hostname and port number** for specific services, helping clients '
    'automatically discover where a service is running without hard-coding addresses.'
)
add_paragraph('Commonly used in:')
add_bullet('VoIP (Voice over IP) services')
add_bullet('SIP (Session Initiation Protocol)')
add_bullet('XMPP (Extensible Messaging)')
add_bullet('Gaming services')
add_bullet('Microsoft Teams / Skype for Business')
add_code_block('_sip._tcp.example.com  10  60  5060  sip.example.com\n    │         │    │    │     └─ Target host\n    │         │    │    └───── Port number\n    │         │    └────────── Weight\n    │         └─────────────── Priority\n    └───────────────────────── Service.Protocol')

# ─────────────────────────────────────────────────────────────
# SECTION 6 — ALIAS RECORDS
# ─────────────────────────────────────────────────────────────
add_heading('Alias Records', 1)

add_paragraph(
    'Alias records are a **Route 53 extension** to standard DNS. They map domain names directly '
    'to AWS resources and solve a key limitation of CNAME records — they work at the apex/root domain.'
)
add_paragraph('Alias records can target:')
add_bullet('Load Balancers (ALB, NLB, CLB)')
add_bullet('CloudFront distributions')
add_bullet('API Gateway endpoints')
add_bullet('S3 Static Websites')
add_bullet('Global Accelerator')
add_bullet('Another Route 53 record in the same hosted zone')

add_paragraph('Features:')
add_bullet('Works at root domain (`example.com`) — unlike CNAME')
add_bullet('Automatically tracks AWS IP changes — no manual updates needed')
add_bullet('No extra DNS query — resolved internally in one step')
add_bullet('Free of charge for queries targeting AWS resources')
add_bullet('Native health check support')

add_code_block('example.com  →  my-alb-1234567890.us-east-1.elb.amazonaws.com  (Alias)')

add_heading('CNAME vs Alias Comparison', 2)
add_table(
    ['Feature', 'CNAME', 'Alias'],
    [
        ['Root Domain Support', 'No — subdomains only', 'Yes — works at apex'],
        ['AWS Native Integration', 'Partial', 'Full native integration'],
        ['Extra DNS Lookup', 'Yes — adds latency', 'No — resolved internally'],
        ['Automatic AWS IP Updates', 'No — manual updates needed', 'Yes — automatic'],
        ['Cost for AWS Targets', 'Standard query charges', 'Free for AWS targets'],
        ['TTL Control', 'Manual', 'Managed by AWS automatically'],
        ['Health Check Support', 'Limited', 'Full support'],
    ]
)

# ─────────────────────────────────────────────────────────────
# SECTION 7 — ROUTING POLICIES
# ─────────────────────────────────────────────────────────────
add_heading('Route 53 Routing Policies', 1)

add_blockquote(
    'Routing Policies define how Route 53 responds to DNS queries. DNS itself does NOT route '
    'traffic like a Load Balancer — it only decides which IP or resource name to return in '
    'the DNS response. The client then connects directly to that resource.'
)

add_heading('Routing Policies Overview', 2)
add_table(
    ['Policy', 'Routes Based On', 'Health Checks', 'Common Use'],
    [
        ['**Simple**', 'Single resource', 'Limited', 'Basic single-server websites'],
        ['**Weighted**', 'Assigned weights', 'Yes', 'A/B testing, gradual rollout'],
        ['**Latency**', 'Lowest network latency', 'Yes', 'Global low-latency applications'],
        ['**Failover**', 'Primary / Secondary health', 'Required', 'Disaster recovery'],
        ['**Geolocation**', 'User country / continent', 'Yes', 'Localization & compliance'],
        ['**Geoproximity**', 'Geography + configurable bias', 'Yes', 'Fine-grained traffic shaping'],
        ['**Multi-Value**', 'Multiple healthy IPs', 'Yes', 'Basic load distribution'],
        ['**IP-Based**', 'Client IP / CIDR range', 'Yes', 'ISP or network-based routing'],
    ]
)

add_heading('Simple Routing', 2)
add_paragraph('Returns a single resource for a domain. The simplest routing policy.')
add_code_block('example.com  →  54.12.34.56')
add_diagram_label('ARCHITECTURE — Simple Routing')
add_code_block(
    'User → DNS Query → Route 53 → Returns 54.12.34.56 → User connects to EC2\n\n'
    'If multiple IPs configured:\n'
    'Route 53 returns ALL IPs → Browser picks one randomly (round-robin)'
)
add_paragraph('Features:')
add_bullet('Simplest routing policy')
add_bullet('Can return multiple IPs — browser chooses randomly')
add_bullet('No health-check based failover')
add_bullet('Best for small or simple single-server applications')
add_paragraph('Use Case: Personal blog hosted on one EC2 instance')

add_heading('Weighted Routing', 2)
add_paragraph('Splits traffic between multiple resources using assigned weights.')
add_code_block(
    'Server A → Weight 80   →  80% of traffic\n'
    'Server B → Weight 20   →  20% of traffic'
)
add_diagram_label('ARCHITECTURE — Weighted Routing')
add_code_block(
    'User Request\n'
    '     │\n'
    '     ▼\n'
    'Route 53 Weighted Policy\n'
    '     │\n'
    '     ├── 80% ──→ Server A (Production v1)\n'
    '     │               IP: 54.12.34.56\n'
    '     │\n'
    '     └── 20% ──→ Server B (New v2 being tested)\n'
    '                     IP: 54.12.34.99'
)
add_paragraph('Features:')
add_bullet('Supports gradual deployments and A/B testing')
add_bullet('Blue-Green deployments — shift traffic incrementally')
add_bullet('Weight range: `0–255`')
add_bullet('Weight `0` stops traffic to a resource without deleting the record')
add_bullet('If all weights are `0`, traffic is distributed equally among all records')
add_paragraph('Use Case: Testing a new application version with a limited percentage of users')

add_heading('Latency Routing', 2)
add_paragraph(
    'Routes users to the AWS region with the **lowest network latency**. This is measured '
    'between the end user\'s network and the AWS region — not purely geographic distance.'
)
add_code_block(
    'India User  → Mumbai Region    (ap-south-1)   ← lowest latency path\n'
    'US User     → Virginia Region  (us-east-1)\n'
    'EU User     → Ireland Region   (eu-west-1)'
)
add_diagram_label('ARCHITECTURE — Latency Routing')
add_code_block(
    'User in India\n'
    '     │\n'
    '     ▼\n'
    'Route 53 Latency Check\n'
    '     │\n'
    '     ├── ap-south-1  (Mumbai)   ← 18ms   ← WINNER\n'
    '     ├── us-east-1   (Virginia) ← 210ms\n'
    '     └── eu-west-1   (Ireland)  ← 120ms\n'
    '     │\n'
    '     ▼\n'
    'User directed to Mumbai EC2 / ALB'
)
add_paragraph('Features:')
add_bullet('Optimizes global application performance')
add_bullet('Based on actual latency measurements, not geographic proximity')
add_bullet('Can integrate with health checks')
add_blockquote('Germany users may be directed to US-East if that region has lower latency than EU regions.')
add_paragraph('Use Case: Global applications requiring fast response times for users worldwide')

add_heading('Failover Routing', 2)
add_paragraph(
    'Routes traffic to a **primary resource** and automatically switches to a **secondary backup** '
    'if the primary becomes unhealthy. Requires health checks.'
)
add_diagram_label('ARCHITECTURE — Failover Routing')
add_code_block(
    'Normal State:\n'
    'User → Route 53 → Primary Server (54.12.34.56) ← Active, Healthy\n'
    '                   Secondary Server (54.99.88.77)  ← Standby\n'
    '\n'
    'Failover State (Primary unhealthy):\n'
    'User → Route 53 → Primary Server (54.12.34.56) ← UNHEALTHY ✗\n'
    '               └─→ Secondary Server (54.99.88.77) ← NOW ACTIVE ✓\n'
    '\n'
    'Route 53 Health Checker detects failure within ~30 seconds\n'
    'DNS TTL determines how quickly clients switch over'
)
add_paragraph('Features:')
add_bullet('Requires health checks on the primary record')
add_bullet('Automatic disaster recovery — no manual intervention')
add_bullet('Active-Passive setup (primary serves all traffic normally)')
add_bullet('Secondary can be in a different region for maximum resilience')
add_paragraph('Use Case: High availability applications with backup disaster recovery regions')

add_heading('Geolocation Routing', 2)
add_paragraph(
    'Routes traffic based on the **geographic location** of the user, determined by their '
    'IP address. Unlike Latency Routing, this is purely location-based regardless of speed.'
)
add_paragraph('Supported levels:')
add_bullet('Continent (e.g. Europe, Asia, North America)')
add_bullet('Country (e.g. Germany, India, United States)')
add_bullet('US State (granular US-only routing)')
add_code_block(
    'Germany Users → German Server  (GDPR compliance)\n'
    'India Users   → India Server   (local language, pricing)\n'
    'Default       → Global Server  (catches all other countries)'
)
add_diagram_label('ARCHITECTURE — Geolocation Routing')
add_code_block(
    'Route 53\n'
    '     │\n'
    '     ├── Country: DE (Germany)  ──→ eu-central-1 (Frankfurt)\n'
    '     ├── Country: IN (India)    ──→ ap-south-1 (Mumbai)\n'
    '     ├── Country: US (United States) ──→ us-east-1 (N. Virginia)\n'
    '     └── Default Record         ──→ us-east-1 (catches everyone else)\n'
    '\n'
    'IMPORTANT: Always configure a Default record.\n'
    'Without it, users from unlisted countries get NXDOMAIN (no DNS response).'
)
add_paragraph('Features:')
add_bullet('Supports localization (language, currency, pricing)')
add_bullet('Useful for legal and compliance requirements (GDPR, data residency)')
add_bullet('Default record strongly recommended to handle unlisted locations')
add_paragraph('Use Case: Region-specific pricing, language, or media content restrictions')

add_heading('Geoproximity Routing', 2)
add_paragraph(
    'Routes traffic based on user location AND resource location, with a configurable **bias** '
    'that expands or shrinks the geographic region served by each resource. '
    'Requires **Route 53 Traffic Flow**.'
)
add_paragraph('Bias values:')
add_bullet('`+1 to +99` → Expands the effective traffic region (attracts more users)')
add_bullet('`-1 to -99` → Shrinks the effective traffic region (repels users to other endpoints)')
add_code_block('US-East Bias +50  →  Receives traffic from a larger geographic area than normal')
add_diagram_label('ARCHITECTURE — Geoproximity Routing with Bias')
add_code_block(
    'Without Bias:\n'
    'US Users  → US-East\n'
    'EU Users  → EU-West\n'
    '(split at geographic midpoint)\n'
    '\n'
    'With US-East Bias +50:\n'
    'US Users  → US-East\n'
    'Most EU Users → US-East  ← bias expands the US-East coverage area\n'
    'Far EU Users  → EU-West\n'
    '\n'
    'Useful when one region has more capacity or lower cost'
)
add_paragraph('Use Case: Intentionally directing more users to a preferred or higher-capacity region')

add_heading('Multi-Value Routing', 2)
add_paragraph(
    'Returns **multiple healthy IP addresses** for a single DNS query. '
    'The client receives up to 8 healthy IPs and picks one — providing basic client-side load distribution.'
)
add_code_block(
    'Healthy:\n'
    '54.12.34.56  ✓\n'
    '54.12.34.99  ✓\n'
    '54.12.34.11  ✓\n'
    '\n'
    'Unhealthy:\n'
    '54.12.34.22  ✗  ← Route 53 will NOT return this IP'
)
add_diagram_label('ARCHITECTURE — Multi-Value Routing')
add_code_block(
    'Route 53 Multi-Value Query Response:\n'
    '[\n'
    '  54.12.34.56,  ← healthy\n'
    '  54.12.34.99,  ← healthy\n'
    '  54.12.34.11,  ← healthy\n'
    '  54.12.34.22   ← EXCLUDED (unhealthy)\n'
    ']\n'
    '\n'
    'Client picks one IP randomly and connects directly\n'
    '\n'
    'NOTE: Multi-Value is NOT a replacement for an Elastic Load Balancer.\n'
    'ELB handles connection-level load balancing, health checks, SSL termination.\n'
    'Multi-Value only controls which IPs the DNS response includes.'
)
add_paragraph('Features:')
add_bullet('Returns up to 8 healthy records per query')
add_bullet('Supports health checks — automatically excludes unhealthy endpoints')
add_bullet('Basic client-side load distribution')
add_bullet('Not a replacement for ELB — DNS does not balance existing connections')
add_paragraph('Use Case: Multiple EC2 instances serving the same application with basic redundancy')

add_heading('IP-Based Routing', 2)
add_paragraph(
    'Routes traffic based on the **client\'s IP address or CIDR range**. '
    'You define CIDR blocks and map each block to a specific endpoint.'
)
add_code_block(
    '203.0.113.0/24   →  Server A  (Corporate office network)\n'
    '198.51.100.0/24  →  Server B  (ISP or branch office network)\n'
    'Everything else  →  Default Server'
)
add_diagram_label('ARCHITECTURE — IP-Based Routing')
add_code_block(
    'Client IP: 203.0.113.50\n'
    '     │\n'
    '     ▼\n'
    'Route 53 CIDR Block Lookup\n'
    '     │\n'
    '     │  203.0.113.50 matches 203.0.113.0/24\n'
    '     ▼\n'
    'Server A  (optimized endpoint for this network)'
)
add_paragraph('Features:')
add_bullet('ISP or network-specific routing for optimized performance')
add_bullet('CIDR-based IP-to-endpoint mapping')
add_bullet('Useful for enterprise branch office routing')
add_bullet('Can direct known bot or crawler IPs to specific handling')
add_paragraph('Use Case: Directing branch offices or specific ISP users to optimized endpoints')

# ─────────────────────────────────────────────────────────────
# SECTION 8 — HEALTH CHECKS
# ─────────────────────────────────────────────────────────────
add_heading('Route 53 Health Checks', 1)

add_paragraph(
    'Route 53 Health Checks continuously monitor application endpoints and automatically stop '
    'routing traffic to unhealthy resources. They are the backbone of Route 53\'s high availability features.'
)
add_blockquote(
    'Think of Health Checks as an automated monitoring and failover system for your infrastructure. '
    'They run 24/7 from multiple AWS locations worldwide.'
)

add_heading('Types of Health Checks', 2)

add_heading('1. Endpoint Health Check', 3)
add_paragraph(
    'Monitors a specific IP address, domain, or URL using **HTTP, HTTPS, or TCP** protocols. '
    'Route 53 sends requests from health checker locations around the world and evaluates the response.'
)
add_code_block('https://example.com/health  →  Expected: HTTP 200 OK')
add_paragraph('You can also configure Route 53 to check for specific text in the response body.')

add_heading('2. Calculated Health Check', 3)
add_paragraph(
    'Combines multiple individual health checks into a single logical result using '
    '**AND, OR, or NOT** conditions. Useful for expressing complex health logic.'
)
add_paragraph('Example:')
add_bullet('Healthy only if **2 out of 3** child checks pass')
add_bullet('Unhealthy if ANY of 5 database checks fails')
add_bullet('Healthy only if BOTH the web tier AND the database tier are healthy')

add_heading('3. CloudWatch Alarm Health Check', 3)
add_paragraph(
    'Uses a **CloudWatch Alarm** state instead of directly checking an endpoint. '
    'The health check is considered healthy when the alarm is in OK state and '
    'unhealthy when the alarm is in ALARM state.'
)
add_blockquote(
    'Route 53 health checkers are public AWS infrastructure and cannot directly reach '
    'resources inside private VPCs. Use CloudWatch Alarm Health Checks for private resources '
    'by publishing custom metrics from your VPC to CloudWatch.'
)

add_heading('Health Check Configuration', 2)
add_table(
    ['Setting', 'Example', 'Notes'],
    [
        ['**Protocol**', 'HTTPS', 'HTTP, HTTPS, or TCP'],
        ['**Port**', '443', 'Standard or custom port'],
        ['**Path**', '`/health`', 'Endpoint that returns 200 when healthy'],
        ['**Check Interval**', '30 sec', '10 sec also available (higher cost)'],
        ['**Failure Threshold**', '3 failed checks', 'Consecutive failures before marking unhealthy'],
        ['**Success Threshold**', '1 success', 'Consecutive successes before marking healthy again'],
        ['**Regions**', 'Multiple', 'Choose which AWS regions run health checks from'],
    ]
)

add_heading('Key Health Check Behavior', 2)
add_bullet('Around **15 global AWS health checkers** monitor endpoints continuously from different regions')
add_bullet('Only `2xx` and `3xx` HTTP responses are considered healthy')
add_bullet('Response time threshold: default 4 seconds for standard, 2 seconds for fast checks')
add_bullet('String matching: Route 53 can check the first 5,120 bytes of the response body')
add_bullet('CloudWatch integration: health check status is published as CloudWatch metrics')
add_bullet('SNS Notifications: trigger alerts via SNS when health status changes')

add_heading('Health Check + Failover Flow', 2)
add_diagram_label('ARCHITECTURE — Health Check Failover')
add_code_block(
    'Route 53 Health Checkers (15 global locations)\n'
    '     │\n'
    '     │  Every 30 seconds: GET https://example.com/health\n'
    '     ▼\n'
    'Primary Server  →  HTTP 200  →  HEALTHY ✓\n'
    '     │\n'
    '     │  [If primary returns errors 3 times in a row]\n'
    '     ▼\n'
    'Primary Server  →  Connection Timeout  →  UNHEALTHY ✗\n'
    '     │\n'
    '     │  Route 53 stops returning Primary IP\n'
    '     ▼\n'
    'Secondary Server  →  Route 53 now returns Secondary IP\n'
    '     │\n'
    '     │  Clients get Secondary IP on next DNS query\n'
    '     ▼\n'
    'Traffic automatically rerouted  →  Failover complete\n'
    '\n'
    'When Primary recovers:\n'
    'Route 53 detects healthy response → Resumes sending traffic to Primary'
)

# ─────────────────────────────────────────────────────────────
# SECTION 9 — TASKS
# ─────────────────────────────────────────────────────────────
add_heading('TASKS — Hands-On Route 53 Labs', 1)

# ── TASK 1 ──
add_heading('Task 1 — Host a Website Using Simple Routing', 2)

add_heading('Objective', 3)
add_bullet('Launch an EC2 instance')
add_bullet('Install a web server')
add_bullet('Configure a custom domain')
add_bullet('Point the domain to EC2 using Route 53')
add_bullet('Access the website publicly')

add_heading('Architecture', 3)
add_diagram_label('ARCHITECTURE — Simple Routing — Task 1')
add_code_block(
    'Internet User\n'
    '     │\n'
    '     ▼\n'
    'www.yourdomain.com\n'
    '     │\n'
    '     ▼  (DNS Query)\n'
    'Route 53 Hosted Zone\n'
    '     │  A Record: yourdomain.com → 54.12.34.56\n'
    '     ▼\n'
    'EC2 Instance (Apache / Nginx)\n'
    'IP: 54.12.34.56  |  Port: 80  |  Public Subnet\n'
    '     │\n'
    '     ▼\n'
    'Response: Hello from AWS Route 53!'
)

add_heading('Step 1 — Purchase a Domain', 3)
add_paragraph('**GoDaddy / Namecheap:**')
add_bullet('Purchase a domain')
add_bullet('Route 53 will manage DNS after configuration')
add_paragraph('**Route 53 (directly):**')
add_bullet('Open AWS Console → Route 53')
add_bullet('Register Domain')
add_bullet('Hosted zone is created automatically after purchase')

add_heading('Step 2 — Launch EC2 Instance', 3)
add_table(
    ['Setting', 'Value'],
    [
        ['AMI', 'Amazon Linux 2023'],
        ['Instance Type', 't2.micro'],
        ['Public IP', 'Enabled'],
        ['Security Group', 'Allow SSH (22) & HTTP (80)'],
    ]
)
add_paragraph('Launch steps:')
add_bullet('Open EC2 → Launch Instance')
add_bullet('Create or select key pair')
add_bullet('Launch instance')
add_bullet('Note the Public IP address')
add_code_block('Example Public IP: 54.12.34.56')

add_heading('Step 3 — Install Apache2 Web Server (Ubuntu)', 3)
add_paragraph('Connect to EC2:')
add_code_block('ssh -i your-key.pem ubuntu@54.12.34.56')
add_paragraph('Install Apache2:')
add_code_block(
    'sudo apt update -y\n'
    'sudo apt install apache2 -y\n\n'
    'sudo systemctl start apache2\n'
    'sudo systemctl enable apache2'
)
add_paragraph('Create sample webpage:')
add_code_block(
    "sudo bash -c 'cat > /var/www/html/index.html << EOF\n"
    '<h1>Hello from AWS Route 53!</h1>\n'
    'EOF\''
)
add_paragraph('Verify in browser:')
add_code_block('http://54.12.34.56')

add_heading('Step 4 — Create Route 53 Hosted Zone', 3)
add_bullet('Open Route 53')
add_bullet('Hosted Zones → Create Hosted Zone')
add_table(
    ['Field', 'Value'],
    [
        ['Domain Name', 'yourdomain.com'],
        ['Type', 'Public Hosted Zone'],
    ]
)
add_paragraph('Copy the 4 generated NS records. Example:')
add_code_block(
    'ns-123.awsdns-45.com\n'
    'ns-678.awsdns-12.net\n'
    'ns-901.awsdns-34.org\n'
    'ns-234.awsdns-56.co.uk'
)

add_heading('Step 5 — Update Nameservers at Registrar', 3)
add_paragraph('**GoDaddy:**')
add_bullet('My Products → DNS')
add_bullet('Nameservers → Change → Custom')
add_bullet('Paste the 4 Route 53 NS records')
add_paragraph('**Namecheap:**')
add_bullet('Domain List → Manage')
add_bullet('Nameservers → Custom DNS')
add_bullet('Paste the 4 Route 53 NS records')
add_paragraph('DNS propagation typically takes a few minutes to up to 48 hours.')

add_heading('Step 6 — Create A Record in Route 53', 3)
add_bullet('Route 53 → Hosted Zone → yourdomain.com')
add_bullet('Create Record')
add_table(
    ['Field', 'Value'],
    [
        ['Record Name', 'www (or leave blank for root)'],
        ['Record Type', 'A'],
        ['Routing Policy', 'Simple'],
        ['Value', 'EC2 Public IP (54.12.34.56)'],
        ['TTL', '300'],
    ]
)

add_heading('Step 7 — Validate DNS Resolution', 3)
add_code_block(
    '# Check DNS propagation\n'
    'nslookup yourdomain.com\n'
    'dig yourdomain.com\n\n'
    '# Expected output\n'
    ';; ANSWER SECTION:\n'
    'yourdomain.com.  300  IN  A  54.12.34.56'
)
add_paragraph('Open in browser:')
add_code_block('http://yourdomain.com')

# ── ROUTING POLICIES TASKS ──
add_heading('Routing Policies — Practice Tasks', 2)

add_heading('Weighted Routing', 3)
add_paragraph('**Purpose:** Split traffic between servers for A/B testing or gradual deployments.')
add_table(
    ['Server', 'Weight', 'Traffic %'],
    [
        ['Server 1 (Production)', '80', '80% of traffic'],
        ['Server 2 (New Version)', '20', '20% of traffic'],
    ]
)
add_paragraph('Configuration:')
add_table(
    ['Field', 'Value'],
    [
        ['Record Type', 'A'],
        ['Routing Policy', 'Weighted'],
        ['Record ID', 'Unique name per record (e.g. server-1, server-2)'],
        ['Weight', '80 or 20'],
    ]
)
add_paragraph('Test weighted distribution:')
add_code_block(
    '# Run 10 DNS lookups and observe IP distribution\n'
    'for i in {1..10}; do\n'
    '  dig +short yourdomain.com\n'
    'done'
)
add_diagram_label('ARCHITECTURE — Weighted Routing Task')
add_code_block(
    'DNS Query → Route 53 Weighted Policy\n'
    '     │\n'
    '     ├── 80% → Server 1 IP (Production)\n'
    '     └── 20% → Server 2 IP (New version being tested)'
)

add_heading('Latency Routing', 3)
add_paragraph('**Purpose:** Route users to the lowest-latency AWS region.')
add_table(
    ['Region Code', 'Location'],
    [
        ['ap-south-1', 'Mumbai, India'],
        ['us-east-1', 'N. Virginia, USA'],
        ['eu-west-1', 'Ireland, EU'],
    ]
)
add_paragraph('Configuration:')
add_table(
    ['Field', 'Value'],
    [
        ['Record Type', 'A'],
        ['Routing Policy', 'Latency'],
        ['Region', 'Select the AWS region where your resource is deployed'],
    ]
)
add_paragraph('Create one Latency record per region pointing to that region\'s server/ALB.')
add_diagram_label('ARCHITECTURE — Latency Routing Task')
add_code_block(
    'User in India → Route 53 measures latency:\n'
    '  ap-south-1:  18ms  ← LOWEST\n'
    '  us-east-1:   210ms\n'
    '  eu-west-1:   120ms\n'
    'Result: User routed to Mumbai server'
)

add_heading('Failover Routing', 3)
add_paragraph('**Purpose:** Disaster recovery with automatic failover to backup server.')
add_paragraph('Step 1 — Create Health Check:')
add_table(
    ['Field', 'Value'],
    [
        ['Protocol', 'HTTP'],
        ['Port', '80'],
        ['Path', '/health'],
        ['Failure Threshold', '3'],
    ]
)
add_paragraph('Step 2 — Primary Record:')
add_table(
    ['Field', 'Value'],
    [
        ['Failover Type', 'Primary'],
        ['Value', 'Primary Server IP'],
        ['Health Check', 'Attach the health check created above'],
    ]
)
add_paragraph('Step 3 — Secondary Record:')
add_table(
    ['Field', 'Value'],
    [
        ['Failover Type', 'Secondary'],
        ['Value', 'Backup Server IP'],
        ['Health Check', 'Optional (recommended)'],
    ]
)
add_paragraph('Test failover — stop the web server on the primary:')
add_code_block(
    '# Stop primary server\n'
    'sudo systemctl stop httpd\n\n'
    '# Watch Route 53 health check go unhealthy (~30-90 seconds)\n'
    '# Then DNS will start returning the secondary IP\n'
    'dig +short yourdomain.com\n\n'
    '# Restore primary\n'
    'sudo systemctl start httpd'
)
add_diagram_label('ARCHITECTURE — Failover Routing Task')
add_code_block(
    'Normal:   Route 53 → Primary (Healthy ✓) → Traffic flows here\n'
    'Failure:  Route 53 → Primary (Unhealthy ✗) → Traffic → Secondary ✓'
)

add_heading('Geolocation Routing', 3)
add_paragraph('**Purpose:** Route traffic based on user\'s country or continent.')
add_table(
    ['Location', 'Server', 'Use Case'],
    [
        ['India', 'India Server IP', 'Local language and pricing'],
        ['USA', 'US Server IP', 'English content, US pricing'],
        ['Default', 'Global Server IP', 'Catch-all for all other countries'],
    ]
)
add_paragraph('Configuration:')
add_table(
    ['Field', 'Value'],
    [
        ['Record Type', 'A'],
        ['Routing Policy', 'Geolocation'],
        ['Location', 'Select Continent, Country, or US State'],
    ]
)
add_blockquote('Always create a Default geolocation record. Without it, users from unlisted countries will get NXDOMAIN errors.')
add_diagram_label('ARCHITECTURE — Geolocation Routing Task')
add_code_block(
    'Route 53 Geolocation\n'
    '  IN (India)  → India EC2\n'
    '  US (USA)    → US EC2\n'
    '  Default     → Global EC2  (catches everyone else)'
)

add_heading('Multi-Value Routing', 3)
add_paragraph('**Purpose:** Return multiple healthy IP addresses for basic client-side load distribution.')
add_table(
    ['Field', 'Value'],
    [
        ['Record Type', 'A'],
        ['Routing Policy', 'Multivalue Answer'],
        ['Health Check', 'Enabled — attach one per IP'],
        ['Value', 'One IP per record (create multiple records)'],
    ]
)
add_paragraph('Create one record per server IP with health checks. Route 53 returns all healthy IPs.')
add_diagram_label('ARCHITECTURE — Multi-Value Routing Task')
add_code_block(
    'Route 53 Multi-Value Response:\n'
    '  54.12.34.56  ✓  (healthy, included in response)\n'
    '  54.12.34.99  ✓  (healthy, included in response)\n'
    '  54.12.34.11  ✗  (unhealthy, excluded from response)\n'
    '\n'
    'Client receives 2 IPs and picks one to connect to'
)

# ── DNS RECORD TYPES EXPLORATION ──
add_heading('Route 53 Record Types — Exploration', 2)
add_paragraph('Explore different Route 53 DNS record types and their use cases:')
add_bullet('**A Record** → Maps domain to IPv4 address')
add_bullet('**AAAA Record** → Maps domain to IPv6 address')
add_bullet('**CNAME Record** → Maps one domain to another domain')
add_bullet('**MX Record** → Configures email routing for the domain')
add_bullet('**TXT Record** → Used for domain verification and email security (SPF, DKIM, DMARC)')
add_bullet('**NS Record** → Defines the authoritative nameservers for the domain')
add_bullet('**SOA Record** → Stores DNS zone administrative information')
add_bullet('**PTR Record** → Used for reverse DNS lookup (IP → domain)')
add_bullet('**Alias Record** → Points root domain to AWS resources like ALB or CloudFront')
add_paragraph(
    'Explore and configure each record type in Route 53 to understand their practical use cases '
    'in real-world AWS architectures.'
)

# ─────────────────────────────────────────────────────────────
# SECTION 10 — KEY TAKEAWAYS
# ─────────────────────────────────────────────────────────────
add_heading('Key Takeaways', 1)

add_bullet('**DNS** is the global system that maps domain names to IP addresses through a hierarchical lookup chain')
add_bullet('**Nameservers** are the authoritative source of DNS records — Route 53 assigns 4 NS per hosted zone')
add_bullet('**Route 53** provides DNS management, intelligent routing, domain registration, and health monitoring in one service')
add_bullet('**Hosted Zones** are containers for DNS records — Public for internet, Private for VPC-internal use')
add_bullet('**Alias Records** are preferred over CNAME for AWS resources — they work at root domain and are free')
add_bullet('**Routing Policies** control which resource IP is returned — Simple, Weighted, Latency, Failover, Geolocation, Geoproximity, Multi-Value, IP-Based')
add_bullet('**Health Checks** continuously monitor endpoints and trigger automatic failover when resources become unhealthy')
add_bullet('**TTL** controls how long DNS responses are cached — lower TTL allows faster updates but increases query volume')
add_bullet('**DNS propagation** takes time — always plan changes in advance and lower TTL before major migrations')
add_bullet('**Route 53 has a 100% availability SLA** — it is the only AWS service with a full uptime guarantee')

add_heading('Explore Routing Policies', 2)
add_bullet('Simple Routing — Single resource, basic use cases')
add_bullet('Weighted Routing — Traffic splitting, A/B testing, Blue-Green deployments')
add_bullet('Latency Routing — Global applications, lowest-latency regional routing')
add_bullet('Failover Routing — Disaster recovery, Active-Passive high availability')
add_bullet('Geolocation Routing — Compliance, localization, region-specific content')
add_bullet('Geoproximity Routing — Fine-grained traffic shaping with bias configuration')
add_bullet('Multi-Value Routing — Basic redundancy across multiple healthy endpoints')
add_bullet('IP-Based Routing — Enterprise networks, ISP-specific or CIDR-based routing')
add_paragraph(
    'Each policy has its own strengths. Real-world architectures often combine multiple policies '
    'using Route 53 Traffic Flow for complex, multi-condition routing logic.'
)

# Closing quote
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), '0F3460')
pPr.append(shd)
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(6)
r = p.add_run('  DNS is the first step of every internet request. Master Route 53 and you control where traffic goes.  ')
set_font(r, size=11, italic=True, color=(230, 230, 255))

doc.save('route53-case-study.docx')
print('Done → route53-case-study.docx')
